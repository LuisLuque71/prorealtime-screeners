"""
Genera el documento de referencia de screeners ProRealTime.
Ejecutar: python generar_doc.py
Salida: ProScreeners_PRT_Referencia.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── helpers ────────────────────────────────────────────────────────────────

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def add_heading(doc, text, level, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
    return h

def add_para(doc, text, bold=False, italic=False, size=10, color=None, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.6)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.8)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10)
        p.add_run(text).font.size = Pt(10)
    else:
        p.add_run(text).font.size = Pt(10)
    return p

def add_score_table(doc, rows):
    """rows = list of (id, pts, condition, explanation)"""
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = tbl.rows[0].cells
    for i, txt in enumerate(["ID", "Pts", "Condición", "Significado"]):
        hdr[i].text = txt
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.size = Pt(9)
        shade_cell(hdr[i], "D9E1F2")
    for (id_, pts, cond, expl) in rows:
        r = tbl.add_row().cells
        r[0].text = id_
        r[1].text = str(pts)
        r[2].text = cond
        r[3].text = expl
        for c in r:
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    # column widths
    for row in tbl.rows:
        row.cells[0].width = Cm(1.0)
        row.cells[1].width = Cm(0.9)
        row.cells[2].width = Cm(7.0)
        row.cells[3].width = Cm(8.0)
    return tbl

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def section_divider(doc):
    doc.add_paragraph()

def add_info_box(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    v = p.add_run(value)
    v.font.size = Pt(10)

def screener_header(doc, name, version, timeframe, universe, objective, edge):
    add_heading(doc, f"{name}  ·  {version}", level=2)
    add_info_box(doc, "Timeframe", timeframe)
    add_info_box(doc, "Universo", universe)
    add_info_box(doc, "Objetivo", objective)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    r = p.add_run("Edge: ")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x37, 0x5C, 0x23)
    v = p.add_run(edge)
    v.font.size = Pt(10)
    v.italic = True

# ─── documento ──────────────────────────────────────────────────────────────

doc = Document()

# márgenes
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.0)

# estilos base
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# ════════════════════════════════════════════════════════════════════
# PORTADA
# ════════════════════════════════════════════════════════════════════
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("PROREALTIME SCREENERS")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run("Manual de Referencia · Diseño, Edge y Protocolo Operativo")
r2.font.size = Pt(14)
r2.italic = True

doc.add_paragraph()
t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run("NYSE · Nasdaq · Diario · v1.0  —  Mayo 2026")
r3.font.size = Pt(11)
r3.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# ÍNDICE CONCEPTUAL
# ════════════════════════════════════════════════════════════════════
add_heading(doc, "Índice y arquitectura del sistema", level=1)
add_para(doc,
    "El sistema completo está dividido en cinco familias, cada una con una filosofía "
    "operativa diferente. Todos los screeners trabajan en timeframe DIARIO sobre NYSE/Nasdaq "
    "y utilizan como mínimo uno de los tres indicadores núcleo: Koncorde (Blai5), "
    "Vigía10 (Blai5) e Ichimoku.", size=10)
doc.add_paragraph()

familias = [
    ("1. Familia Koncorde",
     "Detecta acumulación/distribución institucional comparando NVI (azul = manos fuertes) "
     "con PVI (verde = manos débiles). Cuatro screeners: Acumulación, Primavera, "
     "Acumulación+Primavera combinado y Distribución (bajista)."),
    ("2. Familia Vigía10 / Pre-Breakout",
     "Detecta compresión interna del oscilador SV y sus cruces clave. Tres screeners "
     "orientados a micro-caps y penny stocks: Pre-Breakout v2.0, Absorción Pre-Squeeze v1.5 "
     "y Sniper-Squeeze v5.3."),
    ("3. Familia Ichimoku / Tenkan-Kijun",
     "Screeners basados en la estructura de precio respecto a la nube y la Kijun. "
     "Takana en Kumo (compresión dentro de la nube) e Imán Kijun (reversión a la media)."),
    ("4. Familia Imán Kijun SHORT",
     "Variantes cortas del Imán Kijun: bajo la nube (Imán Kijun SHORT v2.0) y "
     "sobre la nube tras subida parabólica (Imán Kijun SHORT Sobre Nube v1.0)."),
    ("5. Familia Fórmula Omega",
     "Fusión completa de los tres sistemas (Koncorde + Vigía10 + Ichimoku) con doble "
     "temporalidad (CrossDW). El screener más sofisticado del arsenal. "
     "Versiones LARGO y CORTO, objetivo mínimo 20%."),
    ("6. Familia Penny Spring",
     "Penny Spring Extreme v1.1: captura absorción institucional en penny stocks bajo "
     "la nube, con salto brusco del SV como señal de entrada."),
]

for titulo, desc in familias:
    p = doc.add_paragraph()
    r = p.add_run(titulo)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.add_run(f"  —  {desc}").font.size = Pt(10)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 1. FAMILIA KONCORDE
# ════════════════════════════════════════════════════════════════════
add_heading(doc, "1. Familia Koncorde", level=1)
add_para(doc,
    "Koncorde (Blai5, 2008) descompone el volumen en dos índices: NVI (Negative Volume Index, "
    "azul = manos fuertes / institucionales) y PVI (Positive Volume Index, verde = manos "
    "débiles / retail). A esto añade marron (tendencia pura: RSI + MFI + BollOsc + Stoch/3) "
    "y media = EMA(3)(marron). Los cuatro vectores permiten identificar qué manos dominan "
    "el mercado en cada momento.", size=10)

# ── 1.1 Acumulación ──────────────────────────────────────────────
doc.add_paragraph()
screener_header(doc,
    "Koncorde Acumulación + Ichimoku", "v2.1",
    "Diario",
    "$5–$500, volumen medio ≥ $500k/día",
    "Identificar el inicio de la fase de acumulación antes de que el verde confirme la Primavera.",
    "Azul domina sobre verde (institucionales comprando discretamente) con el verde empezando "
    "a girar al alza → transición inminente a Primavera.")

doc.add_paragraph()
add_para(doc, "Condiciones obligatorias (todas deben cumplirse):", bold=True)
gates = [
    "cK1: azul > 0  — smart money activo",
    "cK2: azul > verde  — definición exacta de acumulación: institucionales por delante del retail",
    "cK3: azul > azul[1]  — momentum institucional creciendo sesión a sesión",
    "cK4: verde > verde[1]  — verde sube hoy (el retail empieza a despertar)",
    "cK5: verde > verde[3]  — verde lleva 3 días recuperando (no es ruido)",
    "cV1: filtro > filtro[1]  — Vigía10 mejorando (confirma que el giro tiene fondo)",
]
for g in gates:
    add_bullet(doc, g)

doc.add_paragraph()
add_para(doc, "Salida del screener:", bold=True)
add_para(doc, "Verde, Azul, Marron, Filtro, SV, Kijun, Precio — sin score: cualquier combinación que pase los gates es válida. "
    "Usar Kijun como referencia de soporte y primer target de subida.", indent=True)

add_para(doc, "Ajuste de sensibilidad:", bold=True)
add_para(doc, "Sin señales → bajar minVolDolar a 250k o quitar cK5. "
    "Demasiados → subir minVolDolar a 1M o añadir azul > azul[1] (ya incluido en cK3).", indent=True)

# ── 1.2 Primavera Naciente ───────────────────────────────────────
section_divider(doc)
screener_header(doc,
    "Koncorde Primavera Naciente + Ichimoku", "v2.0",
    "Diario",
    "$5–$500, volumen medio ≥ $1M/día",
    "Capturar el INICIO de la Primavera Koncorde justo cuando el azul cruza cero.",
    "El azul acaba de cruzar sobre cero (institucionales entrando) y el precio ya está sobre "
    "la nube Ichimoku: alta probabilidad de movimiento explosivo inminente.")

doc.add_paragraph()
add_para(doc, "Condiciones obligatorias:", bold=True)
gates2 = [
    "cK1: bsAzul ≤ bsAzulMax (def. 3 barras)  — azul cruzó cero hace máximo 3 días: señal muy fresca",
    "cK2: verde > 0  — Primavera confirmada (ambas manos en positivo)",
    "cK3: marron > media  — tendencia de fondo alcista",
    "cK4: azul > azul[1]  — institucionales siguen comprando",
    "cI1: close > kumoTop  — precio sobre nube: estructura alcista macro",
    "cI2: tenkan > kijun  — impulso de corto plazo positivo",
    "cI3: ssa > ssb  — nube futura alcista (proyección verde en Ichimoku)",
    "cI4: close > close[26]  — chikou libre: sin resistencia en el pasado reciente",
    "cV1: filtro > 0  — Vigía10 positivo (fondo comprador confirmado)",
]
for g in gates2:
    add_bullet(doc, g)

add_para(doc, "Ajuste de sensibilidad:", bold=True)
add_para(doc, "Sin señales → bsAzulMax a 5 o minVolDolar a 500k. "
    "Demasiados → bsAzulMax a 1 o minVolDolar a 2M.", indent=True)

# ── 1.3 Acumulación + Primavera Combinado ────────────────────────
section_divider(doc)
screener_header(doc,
    "Koncorde Acumulación + Primavera Combinado", "v1.0",
    "Diario",
    "$5–$500, volumen medio ≥ $500k/día",
    "Detectar acciones en cualquiera de las dos fases alcistas de Koncorde en un solo pase.",
    "Combina en OR los dos bloques anteriores. Permite supervisar el continuum "
    "Acumulación→Primavera sin ejecutar dos screeners separados.")

doc.add_paragraph()
add_para(doc, "Lógica:", bold=True)
add_para(doc,
    "condicion = cBase AND (acumOK OR primOK)\n"
    "  •  acumOK: azul>0, azul>verde, azul creciendo, verde girando 3 días, filtro mejorando\n"
    "  •  primOK: verde>0, azul>0, verde>marron, marron>media, close>kumoTop, tenkan>kijun, filtro>0 y creciendo\n"
    "Para identificar la fase activa en cada resultado: si close > Nube → Primavera; "
    "si close < Nube → Acumulación.", indent=True)

# ── 1.4 Distribución ─────────────────────────────────────────────
section_divider(doc)
screener_header(doc,
    "Koncorde Distribución", "v1.0",
    "Diario",
    "$5–$500, volumen medio ≥ $3M/día",
    "Detectar distribución institucional activa para posiciones cortas o salidas.",
    "Azul recién cruzado bajo cero (institucionales saliendo) + verde alto (retail aún "
    "muy comprado) + precio perdiendo la nube = trampa completa lista para colapsar.")

doc.add_paragraph()
add_para(doc, "Condiciones obligatorias (todas):", bold=True)
gates3 = [
    "cK1: bsAzul ≤ bsAzulMax (def. 10) AND azul < 0  — salida institucional fresca",
    "cK2: azul < azul[1]  — la salida continúa acelerándose",
    "cK3: verde > 20  — retail MUY comprado: la trampa es real",
    "cK4: verde > marron  — retail comprando en días de alto volumen (oscp > 0)",
    "cK5: marron < marron[1]  — momentum retail cediendo",
    "cI1: close < kumoTop  — precio pierde la nube: soporte estructural roto",
    "cI2: tenkan < kijun  — impulso de corto plazo negativo",
    "cV1: filtro < 0  — Vigía10 negativo",
    "cV2: filtro < filtro[1]  — Vigía10 deteriorándose",
]
for g in gates3:
    add_bullet(doc, g)

add_para(doc, "Columna clave:", bold=True)
add_para(doc, "Trampa = verde – azul: cuanto mayor, más retail está atrapado y mayor es el potencial "
    "de caída cuando comiencen a vender en pánico.", indent=True)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 2. FAMILIA VIGÍA10 / PRE-BREAKOUT
# ════════════════════════════════════════════════════════════════════
add_heading(doc, "2. Familia Vigía10 / Pre-Breakout", level=1)
add_para(doc,
    "Vigía10 (Blai5) es un oscilador de momentum interno compuesto por SV (señal principal), "
    "MR (media de SV), BS/BI (bandas de volatilidad) y filtro = EMA(20) de SV. "
    "Los cruces clave son: SV sobre 0 (Tipo4/TipoB = señal más fuerte) y SV sobre MR "
    "(señal temprana antes del cruce de cero). Las bandas comprimidas indican energía "
    "acumulada pre-squeeze.", size=10)

# ── 2.1 Pre-Breakout v2.0 ────────────────────────────────────────
doc.add_paragraph()
screener_header(doc,
    "Vigía10 Pre-Breakout", "v2.0",
    "Diario",
    "Micro-caps $1–$10, volumen en dólares ≥ $2M/día",
    "Detectar micro-caps en pre-breakout con señales de momentum Vigía10 clasificadas por tipo.",
    "Clasifica señales en cuatro tipos por calidad descendente: TipoB (SV cruza 0 = más "
    "fuerte), TipoA (SV girando cerca de 0), TipoC (todo positivo), TipoD (divergencia "
    "filtro/SV = giro temprano). El screener ordena por tipo primero y score después.")

doc.add_paragraph()
add_para(doc, "Hard gates:", bold=True)
add_bullet(doc, "close ≥ 1 AND ≤ 10, dolarVol ≥ 2M, close > close[1] (precio subiendo hoy)")
add_bullet(doc, "svSubiendo: SV > SV[1] AND SV[1] > SV[2]")
add_bullet(doc, "haySenal: al menos uno de TipoA/B/C/D activo")
add_bullet(doc, "score ≥ 90")

doc.add_paragraph()
add_para(doc, "Tipos de señal (base del scoring — s1):", bold=True)
tipos = [
    ("TipoB", 40, "SV > 0 AND SV[3] ≤ 0 AND svSubiendo",
     "Señal más fuerte: SV cruzó cero hace ≤ 3 barras y sigue subiendo"),
    ("TipoA", 25, "SV < 0 AND SV > -5 AND SV > MR AND svSubiendo",
     "Señal temprana: SV negativo pero sobre MR y girando — pre-cruce de cero"),
    ("TipoC", 15, "SV > 0 AND MR > 0 AND filtro > 0 AND NOT TipoB",
     "Alineación completa alcista: todo positivo, confirmación tardía"),
    ("TipoD", 10, "filtro < 0 AND filtro mejorando AND SV < 0 AND svSubiendo",
     "Divergencia filtro/SV: giro muy temprano, fondo mejora antes que SV"),
]
add_score_table(doc, tipos)

doc.add_paragraph()
add_para(doc, "Scoring adicional (s2–s15):", bold=True)
scores_pb = [
    ("s2", 15, "filtro > 0", "Vigía10 positivo en su tendencia lenta"),
    ("s3", 15, "astro > 0 AND astro[3] ≤ 0", "SV acaba de superar MR (astro = SV - MR)"),
    ("s4", 10, "MR > 0", "Media de SV positiva: presión compradora activa"),
    ("s5", 10, "svConFuerza", "SV acelera: la subida gana velocidad"),
    ("s6", 8,  "astro > 0", "SV por encima de MR"),
    ("s7", 5,  "filtro > filtro[3]", "Filtro mejorando en 3 días"),
    ("s8", 12, "postSqueeze", "Compresión de SV previa: amplitud reciente ≥ 1.2× amplitud anterior"),
    ("s9", 15, "TipoB AND SV[1] ≤ 0", "Cruce de cero justo hoy (frescura máxima del TipoB)"),
    ("s10", 8, "TipoB AND SV[2] ≤ 0", "Cruce hace 1 barra"),
    ("s11", 15, "volRel ≥ 1.5", "Volumen 1.5× la media: interés real"),
    ("s12", 8,  "volRel 1.2–1.5×", "Volumen moderadamente elevado"),
    ("s13", 8,  "close > open", "Vela alcista hoy"),
    ("s14", 12, "close > Highest[10](close[1])", "Ruptura de máximo de 10 días"),
    ("s15", 10, "atrSqueeze: ATR14 < ATR medio × 0.85", "ATR comprimido: energía acumulada en el precio"),
]
add_score_table(doc, scores_pb)

add_para(doc, "Orden de salida:", bold=True)
add_para(doc, "scoreOrden = tipoNum × 1000 + score  → los TipoB (4×1000) aparecen siempre "
    "antes que los TipoA (3×1000) independientemente del score individual.", indent=True)

# ── 2.2 Absorción Pre-Squeeze v1.5 ──────────────────────────────
section_divider(doc)
screener_header(doc,
    "Absorción Pre-Squeeze", "v1.5",
    "Diario",
    "Ultra-penny $0.50–$3.00, avgVol ≥ 200k, volDolar ≥ $300k/día",
    "Detectar absorción institucional climática en pennies con Vigía10 despertando.",
    "Spike de volumen enorme (3–7× media) + precio que NO se mueve tras el spike + "
    "Bollinger y rango de 25 días comprimidos + SV girando al alza = absorción de "
    "oferta antes de un squeeze explosivo.")

doc.add_paragraph()
add_para(doc, "Hard gates:", bold=True)
add_bullet(doc, "close ≥ 0.50 AND ≤ 3.00, avgVol ≥ 200k, volDolar ≥ 300k")
add_bullet(doc, "volX > 3: spike de volumen ≥ 3× la media en las últimas 5 barras")
add_bullet(doc, "rangoPost < 20%: tras el spike el precio no se ha movido más de 20% — absorción real")
add_bullet(doc, "rangoAmp25 < 60%: acción lateral, no en tendencia abierta")
add_bullet(doc, "sv > -50: Vigía10 no en pozo extremo")
add_bullet(doc, "caída diaria < 15%: descarta capitulaciones")
add_bullet(doc, "score ≥ 70")

doc.add_paragraph()
add_para(doc, "Scoring:", bold=True)
scores_abs = [
    ("—", 20, "volX > 5", "Spike muy potente (5× la media)"),
    ("—", 15, "volX > 7", "Spike masivo (7× la media) — absorción extrema"),
    ("—", 20, "rangoPost < 10%", "Precio completamente quieto tras el spike: absorción pura"),
    ("—", 15, "precioNoExploto", "No subió más del 20% desde el mínimo de 5 barras"),
    ("—", 10, "rangoComprimido", "Rango 25 días < 40%: base estrecha"),
    ("—", 10, "bollComprime", "Bollinger comprimido: anchura actual < 85% de su media"),
    ("—", 20, "svGirando", "SV > SV[1] > SV[2]: Vigía10 girando al alza"),
    ("—",  5, "bandaVigia", "Bandas Vigía10 comprimidas (menor 55% del rango de 40 barras)"),
    ("—",  5, "filtroGira", "Filtro Vigía10 mejorando"),
]
add_score_table(doc, scores_abs)

# ── 2.3 Sniper-Squeeze v5.3 ─────────────────────────────────────
section_divider(doc)
screener_header(doc,
    "Vigía10 Sniper-Squeeze", "v5.3",
    "Diario",
    "Penny stocks $0.10–$5, avgVol ≥ 200k, ATR% ≥ 2%",
    "Multi-señal pre-squeeze: compresión de precio + divergencia volumen + absorción + "
    "Takana + squeeze Vigía10 simultáneos.",
    "El screener más complejo de la familia Vigía10: puntúa seis fuentes de compresión "
    "y una señal Vigía activa. Diseñado para penny stocks con Takana Ichimoku como "
    "filtro estructural adicional.")

doc.add_paragraph()
add_para(doc, "Componentes de compresión (núcleo del score):", bold=True)
nucleo_sniper = [
    ("comprimido", 20, "coilPct10 < 30% del rango de 60 días", "Precio muy comprimido vs historia reciente"),
    ("cercaMin",   10, "distMin (vs mínimo 60d) < 25%",          "Cerca del suelo: máximo potencial de rebote"),
    ("divAcum",    25, "precio plano 10d + volumen medio sube",   "Divergencia: precio quieto, dinero entrando"),
    ("absorcion",  20, "ATR bajando + volumen subiendo",          "Absorción: oferta retirada sin mover el precio"),
    ("takLargo",   25, "Takana activa + contexto OK + TK planas", "Fusión Tenkan-Kijun: acumulación Ichimoku"),
    ("bandaMin",   15, "Bandas Vigía10 en mínimo de 10 barras",   "Squeeze Vigía10: energía interna máxima"),
    ("sqVigia",    20, "bandaComp + svDentro + svSube",           "Squeeze clásico de Vigía10"),
]
add_score_table(doc, nucleo_sniper)

add_para(doc, "Señales Vigía10 adicionales (scoring):", bold=True)
sigs_sniper = [
    ("dorada",    20, "SV cruzó MR recientemente AND filtro > 0", "Señal Dorada: cruce SV/MR con fondo positivo"),
    ("niv1",      10, "bsN1 ≤ 5 AND SV > 0",                      "Cruce SV/MR fresco (≤ 5 barras)"),
    ("pre0",       8, "svUp AND mrUp AND SV < MR AND filtro mejora", "Pre-cruce de cero: giro inminente"),
    ("preSV",      6, "SV girando AND SV < 30",                    "SV girando desde valores bajos"),
    ("filtroGira", 8, "filtro > filtro[3] AND > filtro[5]",        "Filtro acelerando al alza"),
    ("filFresco",  8, "filtro cruzó 0 hace ≤ 8 barras",            "Tendencia lenta recién activada"),
]
add_score_table(doc, sigs_sniper)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 3. FAMILIA TAKANA EN KUMO
# ════════════════════════════════════════════════════════════════════
add_heading(doc, "3. Takana en Kumo", level=1)

screener_header(doc,
    "Takana en Kumo", "v6.0",
    "Diario",
    "$1–$500, volumen en dólares ≥ $500k/día",
    "Detectar la fusión Tenkan-Kijun dentro de la nube Ichimoku con SV girando al alza "
    "→ ruptura inminente hacia arriba.",
    "Cuando Tenkan y Kijun convergen dentro de la nube (Takana), el precio está "
    "comprimido entre dos magnetos. Si el SV gira al alza en ese momento, la resolución "
    "estadística es alcista: el breakout va hacia arriba.")

doc.add_paragraph()
add_para(doc, "Hard gates (baseOk — todos obligatorios):", bold=True)
gates_tk = [
    "close ≥ 1 AND ≤ 500, dolarVol ≥ 500k, close > close[1] (precio sube hoy)",
    "enKumo: precio dentro de la nube (kumoBot ≤ close ≤ kumoTop)",
    "kumoFina: grosor de nube < 15% del precio (zona comprimida, no laberinto)",
    "kumoReal: grosor de nube > 1.5% (nube real, no residual)",
    "takanaOk: ≥ 2 de las últimas 4 barras tienen |Tenkan-Kijun|/close < 0.7% (fusión real)",
    "takanaReal: hoy |Tenkan-Kijun|/close ∈ (0.02%, 0.7%) (no exactamente iguales)",
    "takanaAscSuave: punto medio TK hoy > punto medio TK hace 2 barras (ascendente)",
    "svSube: SV > SV[1] > SV[2] (tres barras consecutivas subiendo)",
    "svNoLate: SV < 40 (no overbought)",
    "svNoDeep: SV > -35 (no sobrevendido extremo)",
]
for g in gates_tk:
    add_bullet(doc, g)

doc.add_paragraph()
add_para(doc, "Scoring (score ≥ 50 para entrar):", bold=True)
scores_tk = [
    ("s1",  15, "SV > 0",                              "Impulso positivo confirmado"),
    ("s2",  10, "SV > MR",                             "Momentum activo: SV supera su media"),
    ("s3",  12, "svConFuerza: aceleración de SV",       "El giro gana velocidad"),
    ("s4",  15, "astro cruzó a positivo (SV-MR > 0 hoy, ≤ 0 hace 3 barras)", "Cruce SV/MR muy reciente"),
    ("s5",   8, "filtro > filtro[3]",                  "Tendencia lenta mejorando"),
    ("s6",  10, "MR > 0",                              "Media de SV positiva"),
    ("s7",  12, "filtroDiv: divergencia filtro/SV",    "Giro temprano: filtro mejora antes que SV"),
    ("s8",  20, "kumoAlcista: SSA > SSB",              "Nube alcista: breakout hacia zona libre"),
    ("s9",  10, "kumoGros < 8%",                       "Nube muy fina: mínima resistencia al romper"),
    ("s9b",  8, "SSA > SSA[3]",                        "Nube mejorando en 3 días"),
    ("s10", 15, "volRel ≥ 1.5×",                       "Volumen elevado: interés real detrás del setup"),
    ("s11",  8, "volRel 1.2–1.5×",                     "Volumen moderadamente alto"),
    ("s12",  8, "close > open",                        "Vela alcista hoy"),
    ("s13", 15, "close > Highest[10](close[1])",       "Ruptura del máximo de 10 días"),
    ("s14", 15, "postSqueeze",                         "Compresión previa de SV: energía acumulada"),
]
add_score_table(doc, scores_tk)

add_para(doc, "Salida del screener:", bold=True)
add_para(doc, "score, diffTK0×100 (fusión%), kumoGros (grosor nube%), filtro — "
    "ordenado por score. La fusión% y el grosor de nube son los dos indicadores "
    "de calidad del setup: cuanto más pequeños, mejor.", indent=True)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 4. FAMILIA IMÁN KIJUN (LARGOS Y CORTOS)
# ════════════════════════════════════════════════════════════════════
add_heading(doc, "4. Familia Imán Kijun", level=1)
add_para(doc,
    "La Kijun-sen de Ichimoku es una media móvil especial (punto medio del rango de 26 barras): "
    "cuando está plana actúa como un imán magnético. Los precios alejados de una Kijun "
    "plana tienden estadísticamente a volver a ella. Esta familia captura ese fenómeno "
    "en ambas direcciones.", size=10)

# ── 4.1 Imán Kijun v2.0 ─────────────────────────────────────────
doc.add_paragraph()
screener_header(doc,
    "Imán Kijun", "v2.0  (Largo)",
    "Diario",
    "$0.50–$25, volumen en dólares ≥ $500k/día",
    "Capturar rebotes desde precios muy alejados POR DEBAJO de la Kijun hacia ella.",
    "Precio alejado 15–40% por debajo de una Kijun plana + SV girando al alza + "
    "convergencia activa (distancia al kijun ya se acorta). El kijun plano es el "
    "imán; la distancia es el potencial de recorrido.")

doc.add_paragraph()
add_para(doc, "Hard gates (todos obligatorios):", bold=True)
gates_ik = [
    "close ≥ 0.50 AND ≤ 25, dolarVol ≥ 500k",
    "close < kijun  — precio por debajo del objetivo",
    "distKijun: (kijun-close)/kijun ∈ [15%, 40%]  — alejado pero no en caída libre",
    "cKijunFlat: rango kijun en 10 barras < 8%  — kijun plana (imán real)",
    "cSvGirando: SV > SV[1]  — Vigía10 empieza a girar",
    "cReversionActiva: distKijun < distKijun[1]  — ya se está acercando al kijun hoy",
]
for g in gates_ik:
    add_bullet(doc, g)

doc.add_paragraph()
add_para(doc, "Scoring (score ≥ 70):", bold=True)
scores_ik = [
    ("s1",  20, "kijunRango < 2%",              "Kijun perfectamente plana: imán máximo"),
    ("s2",  12, "kijunRango 2–4%",              "Kijun muy plana"),
    ("s3",   5, "kijunRango 4–8%",              "Kijun moderadamente plana"),
    ("s4",  10, "distKijun 15–20%",             "Potencial mínimo válido"),
    ("s5",  20, "distKijun 20–30%",             "Zona óptima: recorrido/riesgo ideal"),
    ("s6",  15, "distKijun 30–40%",             "Zona extendida: mayor potencial pero más riesgo"),
    ("s7",  20, "distKijun < distKijun[1]",     "Convergencia activa hoy (factor más crítico)"),
    ("s8",  15, "distKijun < distKijun[3]",     "Convergencia sostenida 3 barras"),
    ("s9",  15, "svSubiendo2: SV sube 2 barras consecutivas", "Giro confirmado, no ruido de 1 día"),
    ("s10", 10, "SV > MR",                      "Impulso comprador activo en Vigía10"),
    ("s11", 12, "SV > 0",                       "Señal de giro máxima: SV cruzó a positivo"),
    ("s12", 10, "filtro > filtro[3]",           "Filtro mejorando en 3 días"),
    ("s13",  8, "filtro > -20",                 "Filtro no demasiado hundido: reversión alcanzable"),
    ("s14", 10, "close < tenkan AND tenkan < kijun", "Camino libre sin resistencias medias"),
    ("s15",  8, "tenkan > tenkan[3]",           "Tenkan girando al alza"),
    ("s16", 15, "volRel ≥ 1.5×",               "Volumen elevado"),
    ("s17",  8, "volRel 1.2–1.5×",             "Volumen moderado"),
    ("s18", 10, "close > open",                 "Vela alcista"),
    ("s19",  8, "close > close[1]",             "Precio subiendo"),
    ("s20", 25, "close > nubeTop",              "Sobre la nube: camino libre al kijun"),
    ("s21", 10, "close dentro de la nube",      "Cruzando la nube"),
    ("s22", 12, "close < nubeBot + grosor < 5%","Bajo la nube pero muy fina: obstáculo mínimo"),
    ("s23",  5, "close < nubeBot + grosor 5–15%","Cierta resistencia de nube"),
    ("s24",  8, "ssa0 > ssb0",                  "Nube futura alcista"),
]
add_score_table(doc, scores_ik)

# ── 4.2 Imán Kijun SHORT v2.0 ───────────────────────────────────
section_divider(doc)
screener_header(doc,
    "Imán Kijun SHORT", "v2.0  (Corto — Bajo la Nube)",
    "Diario",
    "$2–$50, volumen en dólares ≥ $1M/día",
    "Capturar el final de un rebote temporal de un stock en downtrend confirmado "
    "(siempre bajo la nube) que se alejó por encima del kijun → vuelta al kijun.",
    "Downtrend confirmado (precio siempre bajo la nube = Ichimoku bajista macro) + "
    "rebote temporal llevó el precio 20–60% sobre el kijun + SV perdiendo fuerza = "
    "la goma estirada hacia arriba vuelve al kijun.")

doc.add_paragraph()
add_para(doc, "Hard gates:", bold=True)
gates_iks = [
    "close ≥ 2 AND ≤ 50, dolarVol ≥ 1M",
    "cBajoNube: close < nubeBot  — downtrend estructural confirmado por Ichimoku",
    "close > kijun (rebote temporal activo)",
    "distAbove: (close-kijun)/kijun ∈ [20%, 60%]  — goma estirada pero no extrema",
    "svBajando: SV < SV[1]  — Vigía10 pierde fuerza",
    "cReversionActiva: distAbove < distAbove[1]  — ya se acerca al kijun hoy",
]
for g in gates_iks:
    add_bullet(doc, g)

doc.add_paragraph()
add_para(doc, "Bloques de scoring principales:", bold=True)
blocks_iks = [
    ("B1", "Rechazo de la Nube (Variante B — edge principal)",
     "s1: precio ayer sobre nubeBot, hoy rechazado (30 pts). "
     "s2: mecha dentro de la nube (25 pts). "
     "s3/s4: rechazo hace 2-3 barras (20/15 pts). "
     "El rechazo explícito de la nube como resistencia es la señal más fiable del setup."),
    ("B2", "Calidad de la estructura bajista",
     "Nube roja actual (SSA<SSB): 12 pts. Nube futura bajista: 8 pts. "
     "Nube gruesa (>10%): 8 pts."),
    ("B3", "Calidad del target (Kijun)",
     "kijunRango < 3%: 20 pts. 3–6%: 12 pts. Kijun bajando: 8 pts."),
    ("B4", "Potencial de recorrido",
     "distAbove 20–30%: 10 pts. 30–45%: 20 pts. 45–60%: 18 pts. "
     "Convergencia 3 barras: 15 pts."),
    ("B5", "Vigía10 — agotamiento del rebote",
     "svBajando2: 15 pts. SV<MR: 10 pts. SV<0: 12 pts. "
     "filtro<filtro[3]: 10 pts. filtro<0: 8 pts."),
    ("B6–8", "Volumen, profundidad y persistencia bajo la nube",
     "Vela bajista con volumen alto: hasta 15 pts. "
     "Distancia a la nube: hasta 26 pts acumulados. "
     "Historial bajo la nube en 5/10/20 barras: hasta 30 pts."),
]
for (bid, btitle, bdesc) in blocks_iks:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    r = p.add_run(f"Bloque {bid} — {btitle}: ")
    r.bold = True
    r.font.size = Pt(10)
    p.add_run(bdesc).font.size = Pt(10)

add_para(doc, "Aviso operativo:", bold=True)
add_para(doc, "Preferentemente para PUT options. Short directo solo en stocks con "
    "dolarVol > $5M para evitar riesgo de squeeze. Stop sugerido: por encima de "
    "la resistencia de nube reciente.", indent=True)

add_para(doc, "Ranking:", bold=True)
add_para(doc, "SCREENER[score × distAbove] — la distancia al kijun domina el orden. "
    "Un setup con 45% de distancia y score=120 supera uno con 22% y score=150 "
    "(5400 vs 3300). Prioriza siempre la goma más estirada.", indent=True)

# ── 4.3 Imán Kijun SHORT Sobre Nube v1.0 ─────────────────────────
section_divider(doc)
screener_header(doc,
    "Imán Kijun SHORT Sobre Nube", "v1.0  (Corto — Sobre la Nube)",
    "Diario",
    "$5–$200, volumen en dólares ≥ $5M/día",
    "Capturar correcciones desde subidas parabólicas en stocks en uptrend "
    "(sobre la nube) hacia el kijun.",
    "Uptrend robusto (sobre la nube ≥ 20 días) + subida parabólica reciente + "
    "stock frenado (sin nuevos máximos) + SV girando a la baja = "
    "corrección al kijun como objetivo natural.")

doc.add_paragraph()
add_para(doc, "Hard gates:", bold=True)
gates_ikn = [
    "close ≥ 5 AND ≤ 200, dolarVol ≥ 5M (liquidez elevada obligatoria)",
    "cSobreNube: close > nubeTop  — uptrend confirmado",
    "kijunLibre: kijun > nubeBot  — kijun no enterrado bajo la nube (camino libre al objetivo)",
    "distAbove: (close-kijun)/kijun ∈ [25%, 70%]  — parabólico pero no extremo",
    "cFrenada: Highest[5] < Highest[15] × 0.98  — el pico quedó atrás, no hace nuevos máximos",
    "svBajando: SV < SV[1]  — Vigía10 perdiendo fuerza",
    "cReversionActiva: distAbove < distAbove[1]  — corrección ya iniciada",
    "close < close[1]  — precio cayendo hoy",
]
for g in gates_ikn:
    add_bullet(doc, g)

doc.add_paragraph()
add_para(doc, "Bloques de scoring (score ≥ 100):", bold=True)
blocks_ikn = [
    ("B1", "Posición del Kijun (calidad del camino al objetivo)",
     "Kijun sobre nubeTop: 25 pts. Kijun dentro de nube: 10 pts. "
     "Kijun ≥5% sobre nubeTop: 10 pts adicionales."),
    ("B2", "Intensidad de la subida parabólica",
     "Subida >80% vs hace 30 barras: 20 pts. 50–80%: 14 pts. 30–50%: 8 pts."),
    ("B3", "Stall — confirmación del techo",
     "Pico hace >1 semana (sin nuevos máximos): 15 pts. Precio bajo Tenkan: 10 pts."),
    ("B4", "Potencial de recorrido",
     "distAbove 25–35%: 10 pts. 35–50%: 20 pts. 50–70%: 16 pts. "
     "Convergencia 3 barras: 15 pts."),
    ("B5", "Volumen decreciente en la cima (distribución)",
     "volMedia5 < volMedia20 × 0.8: 15 pts. ×0.9: 8 pts. "
     "Volumen bajo + vela bajista: 10 pts."),
    ("B6", "Vigía10 — giro bajista",
     "svBajando2: 15 pts. SV<MR: 10 pts. SV<0: 12 pts. filtro<filtro[3]: 10 pts."),
    ("B7–10", "Kijun plano, nube bajista, price action, historial sobre la nube",
     "Kijun plano: hasta 20 pts. Nube actual + futura bajista: 22 pts. "
     "Velas, volumen bajista: hasta 30 pts. Historial sobre nube 5/10/20d: hasta 30 pts."),
]
for (bid, btitle, bdesc) in blocks_ikn:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    r = p.add_run(f"Bloque {bid} — {btitle}: ")
    r.bold = True
    r.font.size = Pt(10)
    p.add_run(bdesc).font.size = Pt(10)

add_para(doc, "Aviso operativo:", bold=True)
add_para(doc, "Probabilidad MENOR que el SHORT bajo la nube (es una corrección dentro de "
    "un uptrend, no una inversión). Usar para PUT options de 4–8 semanas. "
    "Ranking = score × distAbove (la distancia manda).", indent=True)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 5. FÓRMULA OMEGA
# ════════════════════════════════════════════════════════════════════
add_heading(doc, "5. Fórmula Omega — Fusión Total", level=1)
add_para(doc,
    "La Fórmula Omega es el screener más completo del sistema. Fusiona los tres indicadores "
    "núcleo (Koncorde, Vigía10 e Ichimoku) con validación de doble temporalidad "
    "diaria+semanal (CrossDW de Blai5) y un sistema de scoring de 5 bloques. "
    "Objetivo: capturar movimientos del 20%+ con alta probabilidad.", size=10)

doc.add_paragraph()
add_para(doc, "Arquitectura CrossDW (doble temporalidad):", bold=True)
add_para(doc,
    "Dado que ProScreeners trabajan en timeframe diario, la escala semanal se aproxima "
    "multiplicando todos los periodos × 5: EMA[3]→[15], RSI[14]→[70], MFI[14]→[70], "
    "Bollinger[25]→[125], Stoch[21,3]→[105,15], Highest[90]→[450]. "
    "El CrossDW detecta cuando el marronW (Koncorde semanal) cruza su mediaW "
    "(EMA semanal) en la dirección correcta Y ese cruce es reciente (≤ 25 días diarios "
    "≈ 5 semanas). Esto elimina tendencias crónicas que no tienen momentum fresco.",
    indent=True)

# ── 5.1 Omega LARGO ──────────────────────────────────────────────
doc.add_paragraph()
screener_header(doc,
    "Fórmula Omega LARGO", "v1.0",
    "Diario",
    "$5–$200, volumen en dólares ≥ $1M/día",
    "Identificar setups alcistas de alta probabilidad con objetivo ≥ 20% de subida.",
    "Cuando el azul (ballenas/NVI) cruza sobre cero Y el SV de Vigía10 cruza sobre "
    "cero o sobre MR, los institucionales están entrando mientras el momentum interno "
    "se activa. El CrossDW confirma que hay gasolina semanal. Score ≥ 200 asegura "
    "alineación de calidad en múltiples frentes.")

doc.add_paragraph()
add_para(doc, "Puertas núcleo (no negociables — deben cumplirse las TRES):", bold=True)
gates_ol = [
    "hayBallena: sK1 > 0 (azul cruzó 0 hace ≤ 5 días) OR sK2 > 0 (azul positivo y creciendo 2 días)",
    "hayMomentum: sV1 > 0 (SV cruzó 0 hace ≤ 3 días) OR sV2 > 0 (SV cruzó MR hace ≤ 5 días)",
    "crossDW: bsCrossDW ≤ 25 AND marronW > mediaW  — respaldo semanal alcista fresco",
    "cBase: volDolar ≥ 1M, precio ∈ [5, 200]",
    "score ≥ 200",
]
for g in gates_ol:
    add_bullet(doc, g)

doc.add_paragraph()
add_para(doc, "Scoring por bloques (máximo teórico: 360 pts):", bold=True)

# Bloque K
add_para(doc, "Bloque K — Koncorde Alcista (máx. 103 pts)", bold=True, color=(0x00, 0x4E, 0x96))
scores_ok = [
    ("K1", 35, "bsAzul ≤ 5",                          "NÚCLEO: azul cruzó 0 hace ≤ 5 días — ballenas acaban de entrar"),
    ("K2", 20, "azul > 0 AND azul > azul[1] AND azul > azul[2]", "Entrada institucional sostenida 2 sesiones"),
    ("K3", 15, "azul > verde AND azul > 0",            "Institucionales antes que el retail (acumulación pre-pública)"),
    ("K4", 15, "verde < 0 AND azul > 0",               "Retail vendiendo, ballenas comprando — entrada temprana de calidad"),
    ("K5", 10, "marron > media AND marron > marron[1]","Momentum de fondo acelerando"),
    ("K6",  8, "verde > 0 AND azul > 0",               "Spring: toda la ola comprando (confirmación tardía)"),
]
add_score_table(doc, scores_ok)

# Bloque V
doc.add_paragraph()
add_para(doc, "Bloque V — Vigía10 Alcista (máx. 102 pts)", bold=True, color=(0x37, 0x5C, 0x23))
scores_ov = [
    ("V1", 30, "bsSV0 ≤ 3 AND SV > 0",              "SEÑAL TIPO4: SV cruzó 0 hace ≤ 3 días — la más fuerte"),
    ("V2", 20, "bsSVMR ≤ 5",                          "SV cruzó MR hace ≤ 5 días: señal temprana pre-cruce de 0"),
    ("V3", 15, "SV > 0 AND SV > SV[1] > SV[2]",      "Momentum confirmado: SV positivo y acelerando"),
    ("V4", 15, "bsFiltro ≤ 8",                        "Filtro cruzó 0 hace ≤ 8 días: tendencia lenta se activa"),
    ("V5", 12, "SV > MR > 0 AND filtro > 0",          "Alineación perfecta Vigía10: todo positivo"),
    ("V6", 10, "filtro > filtro[1] AND filtro > filtro[3]", "Inercia lenta acompañando"),
]
add_score_table(doc, scores_ov)

# Bloque I
doc.add_paragraph()
add_para(doc, "Bloque I — Ichimoku Alcista (máx. 39 pts)", bold=True, color=(0x7B, 0x36, 0x0A))
scores_oi = [
    ("I1", 15, "close > kumoTop",           "Precio sobre la nube: estructura macro alcista"),
    ("I2", 10, "tenkan > kijun",            "Impulso de corto plazo positivo"),
    ("I3",  8, "ssa > ssb",                 "Nube futura alcista: proyección positiva"),
    ("I4",  6, "close > kijun AND close < kumoTop", "Pre-ruptura de nube: oportunidad antes de la confirmación total"),
]
add_score_table(doc, scores_oi)

# Bloque C
doc.add_paragraph()
add_para(doc, "Bloque C — Compresión/Energía Acumulada (máx. 30 pts)", bold=True)
scores_oc = [
    ("C1", 12, "ATR14 < ATR medio × 0.85",              "ATR comprimido: energía cinética acumulada en el precio"),
    ("C2", 10, "anchoBandas < anchMedia × 0.80",         "Bandas Vigía10 comprimidas: squeeze interno del oscilador"),
    ("C3",  8, "SV dentro de bandas AND svSubiendo",     "Compresión interna activa con SV girando"),
]
add_score_table(doc, scores_oc)

# Bloque P
doc.add_paragraph()
add_para(doc, "Bloque P — Precio/Volumen (máx. 36 pts)", bold=True)
scores_op = [
    ("P1", 12, "volRel ≥ 1.5×",                         "Volumen significativo"),
    ("P2",  6, "volRel 1.2–1.5×",                       "Volumen moderado"),
    ("P3",  8, "close > open",                          "Vela alcista"),
    ("P4", 10, "close > Highest[10](Close[1])",         "Ruptura de máximo de 10 días"),
]
add_score_table(doc, scores_op)

# Bloque DW
doc.add_paragraph()
add_para(doc, "Bloque DW — CrossDW Semanal (máx. 50 pts)", bold=True, color=(0x70, 0x30, 0xA0))
scores_dw = [
    ("DW1", 25, "bsCrossDW ≤ 5",                        "Cruce DW esta semana o la anterior: frescura máxima"),
    ("DW2", 15, "bsCrossDW ≤ 15 AND azulW > 0",         "Cruce DW reciente + ballenas también en semanal"),
    ("DW3", 10, "azulW > 0 AND azulW > azulW[5]",       "Azul semanal positivo y creciendo: institucionales en semanal"),
]
add_score_table(doc, scores_dw)

doc.add_paragraph()
add_para(doc, "Parámetros ajustables:", bold=True)
params_ol = [
    "scoreMin = 200  →  pocas señales: bajar a 185 / muchas: subir a 215",
    "bsAzulMax = 5  →  ventana del cruce azul (días)",
    "bsCrossDWMax = 25  →  ventana CrossDW en días diarios ≈ 5 semanas semanal",
    "minVolDolar = 1M, minPrice = 5, maxPrice = 200",
]
for p in params_ol:
    add_bullet(doc, p)

# ── 5.2 Omega CORTO ─────────────────────────────────────────────
section_divider(doc)
screener_header(doc,
    "Fórmula Omega CORTO", "v1.0",
    "Diario",
    "$5–$500, volumen en dólares ≥ $3M/día",
    "Identificar setups bajistas de alta probabilidad con objetivo ≥ 20% de caída.",
    "El azul (NVI) cruza bajo cero: institucionales saliendo mientras el retail "
    "sigue comprado (verde alto = trampa). El SV gira a la baja. El CrossDW confirma "
    "que el giro semanal también está en marcha. La trampa (verde−azul) cuantifica "
    "el potencial de caída.")

doc.add_paragraph()
add_para(doc, "Diferencias clave respecto a LARGO:", bold=True)
diffs = [
    "minVolDolar = 3M (mayor liquidez: préstamo de acciones más fácil)",
    "maxPrice = 500 (permite grandes caps en distribución)",
    "bsAzulMax = 8 (ventana algo mayor para el cruce bajista)",
    "scoreMin = 210",
    "CrossDW bajista: marronW CROSSES UNDER mediaW + marronW < mediaW",
    "Columna extra Trampa = verde − azul: cuanto mayor, más retail atrapado",
    "Sin Bloque C (compresión): los cortos no necesitan compresión previa",
    "I3 bajista: ssb > ssa (nube roja). I4: close < kijun AND close > kumoBottom",
    "P signals: close < open, close < Lowest[10](Close[1])",
]
for d in diffs:
    add_bullet(doc, d)

doc.add_paragraph()
add_para(doc, "Puertas núcleo:", bold=True)
add_bullet(doc, "hayBallena: sK1 > 0 (azul cruzó 0 a la baja) OR sK2 > 0 (azul negativo cayendo 2 días)")
add_bullet(doc, "hayMomentum: sV1 > 0 (SV cruzó 0 a la baja) OR sV2 > 0 (SV cruzó bajo MR)")
add_bullet(doc, "crossDW: bsCrossDW ≤ 25 AND marronW < mediaW")

doc.add_paragraph()
add_para(doc, "Scoring Koncorde Bajista — diferencias clave:", bold=True)
scores_ok2 = [
    ("K1", 35, "bsAzul ≤ 8 AND azul < 0",              "NÚCLEO: azul cruzó 0 bajista hace ≤ 8 días"),
    ("K2", 20, "azul < 0 AND azul < azul[1] AND azul < azul[2]", "Salida institucional sostenida"),
    ("K3", 15, "trampa > 20 AND azul < 0",             "Retail atrapado: verde muy por encima del azul negativo"),
    ("K4", 15, "verde > 20 AND azul < 0",              "Verde muy positivo: muchas manos que venderán en pánico"),
    ("K5", 10, "marron < media AND marron < marron[1]","Momentum general girándose"),
    ("K6", 12, "verde < 0 AND azul < 0",               "Pánik: todo el mercado vendiendo — aceleración bajista"),
]
add_score_table(doc, scores_ok2)

add_para(doc, "Notas operativas:", bold=True)
add_para(doc, "• Usar solo en acciones con borrowable shares disponibles.\n"
    "• Confirmar en gráfico semanal antes de entrar.\n"
    "• Stop sugerido: por encima de la Tenkan diaria.\n"
    "• Target mínimo: 20% desde entrada (primer objetivo: Kijun semanal).", indent=True)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 6. PENNY SPRING EXTREME
# ════════════════════════════════════════════════════════════════════
add_heading(doc, "6. Penny Spring Extreme", level=1)

screener_header(doc,
    "Penny Spring Extreme", "v1.1",
    "Diario",
    "Penny stocks $0.10–$5, volumen en dólares ≥ $300k/día",
    "Detectar absorción institucional en pennies hundidas bajo la nube con salto "
    "brusco del SV como señal de compra institucional súbita.",
    "Penny bajo la nube (castigada o lateral largo) + volumen climático (≥3×) "
    "con precio quieto (absorción pura) + SV saltando bruscamente desde su mínimo reciente. "
    "Target natural: EMA21, Kijun, kumoBot.")

doc.add_paragraph()
add_para(doc, "Tres pilares obligatorios (todos en hard gate):", bold=True)
pilares = [
    ("Pilar 1 — Posición bajo la nube",
     "bajaNube: close < kumoBot (HARD GATE). La nube es la resistencia. "
     "El precio hundido bajo ella = stock castigado con soporte potencial en la "
     "nube como primer objetivo."),
    ("Pilar 2 — Absorción institucional",
     "climVol: al menos 1 de las últimas 5 barras tiene volumen ≥ 3× la media. "
     "absorcionQuieta OR divAcum: el precio NO se ha movido (rango <10%) a pesar "
     "del spike de volumen = alguien absorbió toda la oferta sin levantar el precio."),
    ("Pilar 3 — Vigía10 girando",
     "svSube OR svSpiked: SV subiendo 2 barras consecutivas, "
     "o salto brusco (svJump ≥ 25 desde el mínimo de 5 barras = señal MNDR)."),
]
for (titulo, desc) in pilares:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    r = p.add_run(f"{titulo}: ")
    r.bold = True
    r.font.size = Pt(10)
    p.add_run(desc).font.size = Pt(10)

doc.add_paragraph()
add_para(doc, "Scoring clave (score ≥ 50):", bold=True)
scores_pse = [
    ("s6",  25, "absorcionQuieta (vol ≥3× + rango <10%)",  "SEÑAL CORE: absorción pura de oferta"),
    ("s7",  15, "divAcum (precio plano 10d + vol medio sube)", "Acumulación acumulativa confirmada"),
    ("s20", 20, "svJump ≥ 40",                             "Salto masivo de SV: compra institucional súbita (MNDR edge)"),
    ("s21", 12, "svJump 25–40",                            "Salto significativo de SV"),
    ("s13", 20, "svDesdePozo: SV[5] < -40 AND SV > SV[5]","SV remontando desde el peor momento"),
    ("s15", 20, "svCruzaPos: SV cruzó 0 hace ≤ 3 barras", "Señal máxima: cruce de cero de SV"),
    ("s1",  20, "distMin60 < 10%",                        "En mínimos absolutos: máximo potencial de rebote"),
    ("s4",  15, "distNube < 15%",                         "Muy cerca del borde inferior de la nube: pull rápido posible"),
    ("s8",  15, "spikeX ≥ 5",                             "Spike muy potente (5× la media)"),
]
add_score_table(doc, scores_pse)

doc.add_paragraph()
add_para(doc, "Edge visual del MNDR:", bold=True)
add_para(doc,
    "La señal más característica de este screener es el 'salto brusco del SV' "
    "(svJump = SV actual − mínimo de SV en 5 barras ≥ 25). En el gráfico de Vigía10 "
    "se ve como una vela larga vertical ascendente desde el suelo. "
    "Esto refleja que en una sola sesión entró suficiente dinero institucional para "
    "mover el NVI/momentum interno aunque el precio apenas se movió.", indent=True)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# RESUMEN COMPARATIVO
# ════════════════════════════════════════════════════════════════════
add_heading(doc, "Tabla resumen del sistema", level=1)

tbl = doc.add_table(rows=1, cols=6)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
hdr = tbl.rows[0].cells
for i, txt in enumerate(["Screener", "Dirección", "Precio", "Vol mín $", "Señal núcleo", "Score mín"]):
    hdr[i].text = txt
    hdr[i].paragraphs[0].runs[0].bold = True
    hdr[i].paragraphs[0].runs[0].font.size = Pt(8)
    shade_cell(hdr[i], "1F497D")
    hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

rows_res = [
    ("Koncorde Acumulación v2.1", "Largo", "5–500", "500k", "Azul>verde, verde girando", "—"),
    ("Koncorde Primavera Naciente v2.0", "Largo", "5–500", "1M", "bsAzul ≤ 3 + Ichimoku", "—"),
    ("Koncorde Acum+Prim Combinado v1.0", "Largo", "5–500", "500k", "acumOK OR primOK", "—"),
    ("Koncorde Distribución v1.0", "Corto/Salida", "5–500", "3M", "Azul recién <0 + verde>20", "—"),
    ("Vigía10 Pre-Breakout v2.0", "Largo", "1–10", "2M", "TipoB/A/C/D + svSubiendo", "90"),
    ("Absorción Pre-Squeeze v1.5", "Largo", "0.50–3", "300k", "volX>3 + rangoPost<20%", "70"),
    ("Vigia10 Sniper-Squeeze v5.3", "Largo", "0.10–5", "200k avgVol", "nucleo>0 + sigVigia", "50"),
    ("Takana en Kumo v6.0", "Largo", "1–500", "500k", "Takana + enKumo + svSube", "50"),
    ("Imán Kijun v2.0", "Largo", "0.50–25", "500k", "dist 15–40% + kijunFlat + svGira", "70"),
    ("Imán Kijun SHORT v2.0", "Corto", "2–50", "1M", "cBajoNube + goma 20–60%", "65"),
    ("Imán Kijun SHORT Sobre Nube v1.0", "Corto", "5–200", "5M", "cSobreNube + parabólico", "100"),
    ("Fórmula Omega LARGO v1.0", "Largo", "5–200", "1M", "hayBallena+hayMomentum+crossDW", "200"),
    ("Fórmula Omega CORTO v1.0", "Corto", "5–500", "3M", "hayBallena+hayMomentum+crossDW", "210"),
    ("Penny Spring Extreme v1.1", "Largo", "0.10–5", "300k", "bajaNube + climVol + absorción + svSube/Spike", "50"),
]

for row_data in rows_res:
    r = tbl.add_row().cells
    for i, val in enumerate(row_data):
        r[i].text = val
        for p in r[i].paragraphs:
            for run in p.runs:
                run.font.size = Pt(8)

# columnas
for row in tbl.rows:
    row.cells[0].width = Cm(5.2)
    row.cells[1].width = Cm(1.8)
    row.cells[2].width = Cm(1.8)
    row.cells[3].width = Cm(1.4)
    row.cells[4].width = Cm(5.2)
    row.cells[5].width = Cm(1.6)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════
# GLOSARIO
# ════════════════════════════════════════════════════════════════════
add_heading(doc, "Glosario de términos", level=1)

glosario = [
    ("Azul (Koncorde)", "NVI normalizado. Representa las manos fuertes/institucionales. Azul > 0 = institucionales comprando."),
    ("Verde (Koncorde)", "PVI + tendencia (oscp + marron). Representa las manos débiles/retail. Verde > azul = retail dominando."),
    ("Marron (Koncorde)", "Tendencia pura: (RSI[14] + MFI[14] + BollOsc[25] + Stoch[21,3]/3) / 2."),
    ("Media (Koncorde)", "EMA(3) del marron: tendencia de muy corto plazo."),
    ("Primavera (Koncorde)", "Fase alcista: azul > 0 Y verde > 0. Ambas manos comprando."),
    ("Acumulación (Koncorde)", "Fase previa a Primavera: azul > 0, verde < 0. Institucionales comprando, retail vendiendo."),
    ("Distribución (Koncorde)", "Fase bajista: azul < 0, verde > 0. Institucionales saliendo, retail atrapado."),
    ("Pánik (Koncorde)", "Fase de capitulación: azul < 0, verde < 0. Todo el mercado vendiendo."),
    ("Trampa (Cortos)", "verde − azul. Cuantifica cuánto retail está atrapado en una distribución. Mayor = más potencial bajista."),
    ("SV (Vigía10)", "Señal principal del oscilador interno. Cruce de 0 = Tipo4/TipoB (señal más fuerte)."),
    ("MR (Vigía10)", "Media de SV. Cruce SV sobre MR = señal temprana antes del cruce de cero."),
    ("Filtro (Vigía10)", "EMA(20) del SV. Tendencia lenta del momentum. filtro > 0 confirma estructura positiva."),
    ("BS / BI (Vigía10)", "Banda Superior / Inferior. Su compresión = squeeze = energía acumulada pre-movimiento."),
    ("Tipo4 / TipoB", "SV cruza 0 hacia arriba: señal más fuerte de Vigía10. Equivalente al azul cruzando 0 en Koncorde."),
    ("Takana", "Fusión de Tenkan y Kijun: |Tenkan − Kijun| < tolerancia. Zona de máxima compresión Ichimoku."),
    ("CrossDW (Blai5)", "Cruce de doble temporalidad D+W: marronW cruzando su mediaW. Detecta el giro semanal en datos diarios."),
    ("barssince()", "Función PRT: cuenta barras desde que una condición fue verdadera. Permite detectar frescura de señales."),
    ("bsAzul", "Barras desde que el azul cruzó cero. bsAzul ≤ 5 = cruce muy fresco."),
    ("bsCrossDW", "Barras desde que el marronW cruzó su mediaW. bsCrossDW ≤ 25 ≈ cruce en últimas 5 semanas."),
    ("kumoTop / kumoBot", "Techo y suelo de la nube Ichimoku (desplazada 26 barras). Max/Min(SSA, SSB)."),
    ("Kijun-sen", "Media del rango de 26 barras de Ichimoku. Cuando está plana actúa como imán estadístico."),
    ("Tenkan-sen", "Media del rango de 9 barras. Impulso de corto plazo. Tenkan > Kijun = momentum positivo."),
    ("SSA / SSB", "Senkou Span A y B: forman la nube. SSA = (Tenkan+Kijun)/2 desplazado 26 barras. SSB = media rango 52 barras."),
    ("distKijun / distAbove", "Distancia porcentual del precio al Kijun. La métrica principal del Imán Kijun."),
    ("svJump", "Salto del SV respecto a su mínimo de 5 barras. svJump ≥ 25 = señal MNDR de compra institucional súbita."),
    ("astro", "SV − MR. Cuando astro cruza de negativo a positivo: cruce SV/MR fresco."),
    ("postSqueeze", "Amplitud de SV en últimas 5 barras ≥ 1.2× amplitud en las 15 barras anteriores: explosión tras compresión."),
    ("absorcionQuieta", "Volumen ≥ 3× la media en una barra con rango < 10%: alguien absorbe oferta sin mover el precio."),
    ("TotalPrice (PRT)", "Precio típico (High+Low+Close)/3. Base de cálculo de MFI, RSI y Bollinger en Koncorde."),
]

for (term, defi) in glosario:
    p = doc.add_paragraph()
    r = p.add_run(f"{term}: ")
    r.bold = True
    r.font.size = Pt(9)
    p.add_run(defi).font.size = Pt(9)

# ─── guardar ────────────────────────────────────────────────────────────────
outpath = r"c:\Users\lluqjap\VSCODE\prorealtime-screeners\ProScreeners_PRT_Referencia.docx"
doc.save(outpath)
print(f"Documento guardado en: {outpath}")
