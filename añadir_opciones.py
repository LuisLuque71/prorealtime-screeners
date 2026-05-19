"""
Añade el capítulo de estrategias con opciones al documento existente.
Ejecutar: python3 añadir_opciones.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── helpers (copiar del script anterior) ───────────────────────────────────

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
    return h

def add_para(doc, text, bold=False, italic=False, size=10, color=None, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.6)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_bullet(doc, text, bold_prefix=None, indent_cm=0.8):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(indent_cm)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10)
        p.add_run(text).font.size = Pt(10)
    else:
        p.add_run(text).font.size = Pt(10)
    return p

def add_info_box(doc, label, value, label_color=(0x1F, 0x49, 0x7D)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(*label_color)
    v = p.add_run(value)
    v.font.size = Pt(10)

def section_divider(doc):
    doc.add_paragraph()

def add_strategy_table(doc, rows):
    """rows = list of (campo, valor)"""
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (campo, valor) in enumerate(rows):
        row = tbl.rows[i].cells
        row[0].text = campo
        row[0].paragraphs[0].runs[0].bold = True
        row[0].paragraphs[0].runs[0].font.size = Pt(9)
        shade_cell(row[0], "E9EFF8")
        row[1].text = valor
        for p in row[1].paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
    for row in tbl.rows:
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(12.5)
    return tbl

def add_legs_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.size = Pt(9)
        shade_cell(hdr[i], "2E5090")
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row_data in rows:
        r = tbl.add_row().cells
        for i, val in enumerate(row_data):
            r[i].text = val
            for p in r[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return tbl

def screener_opt_header(doc, name, direccion, color_dir):
    p = doc.add_heading(name, level=3)
    p.paragraph_format.space_before = Pt(12)
    # dirección badge
    pb = doc.add_paragraph()
    r = pb.add_run(f"  [{direccion}]  ")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(*color_dir)

# ─── constantes de color ─────────────────────────────────────────────────────
LARGO_COLOR  = (0x00, 0x70, 0x00)   # verde oscuro
CORTO_COLOR  = (0xC0, 0x00, 0x00)   # rojo oscuro
MIXED_COLOR  = (0x70, 0x30, 0xA0)   # morado

# ─── cargar documento ────────────────────────────────────────────────────────
docpath = r"c:\Users\lluqjap\VSCODE\prorealtime-screeners\ProScreeners_PRT_Referencia.docx"
doc = Document(docpath)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# CAPÍTULO: ESTRATEGIAS CON OPCIONES
# ════════════════════════════════════════════════════════════════════
add_heading(doc, "7. Estrategias con Opciones por Screener", level=1)

add_para(doc,
    "Este capítulo asigna a cada screener la estrategia de opciones más eficiente "
    "en términos de coste/beneficio, teniendo en cuenta el perfil de volatilidad implícita "
    "(IV) típico del universo de acciones que detecta, el movimiento esperado y el "
    "horizonte temporal del setup.", size=10)

doc.add_paragraph()

# ─── Marco conceptual ────────────────────────────────────────────────────────
add_heading(doc, "Marco conceptual: IV, tiempo y estructura del trade", level=2)

add_para(doc,
    "Las tres variables que determinan qué estrategia usar son:", bold=True)

conceptos = [
    ("IV Implícita (Implied Volatility)",
     "Mide cuánto 'cuesta' el tiempo en las opciones. IV alta = opciones caras → "
     "mejor vender premium o usar spreads para reducir el coste. "
     "IV baja = opciones baratas → mejor comprar directamente (long call / long put). "
     "En los screeners de compresión (Takana, Acumulación), el precio lleva días quieto → "
     "IV generalmente baja → momento ÓPTIMO para comprar opciones."),
    ("Horizonte temporal del setup",
     "Cada screener tiene un horizonte diferente. La regla general: "
     "vencimiento = 2 a 2.5× el tiempo estimado al target. "
     "Esto da margen para el movimiento sin pagar theta excesivo en las últimas 2 semanas."),
    ("Delta y selección de strike",
     "ATM (delta ~0.50): máxima sensibilidad al movimiento del precio, equilibrio coste/beneficio. "
     "OTM 1 strike (delta ~0.30-0.35): más barato, necesita mayor movimiento para rentabilizar. "
     "ITM (delta ~0.70-0.80): actúa casi como el subyacente, menor riesgo de tiempo, más caro. "
     "Regla práctica: para setups de 20%+ usar ATM o 1 strike OTM."),
]

for (titulo, desc) in conceptos:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    r = p.add_run(f"{titulo}: ")
    r.bold = True
    r.font.size = Pt(10)
    p.add_run(desc).font.size = Pt(10)

doc.add_paragraph()
add_para(doc, "Clasificación de IV por tipo de universo:", bold=True)

iv_tabla = [
    ("Tipo de acción", "IV típica", "Estrategia recomendada", "Estrategia alternativa (IV alta)"),
    ("Large cap $50+, vol >$10M", "20–35%", "Long Call / Long Put", "Bull/Bear Call-Put Spread"),
    ("Mid cap $10–$50, vol $2M–$10M", "30–55%", "Spread vertical o Diagonal", "Credit Spread"),
    ("Small cap $5–$15, vol $1M–$3M", "50–80%", "Spread vertical",             "Diagonal o Poor Man's"),
    ("Micro/penny $1–$5, vol <$2M",   "80–150%","Spread muy ajustado",          "Evitar largo de prima pura"),
    ("Ultra-penny $0.10–$1",           ">150%",  "Evitar opciones / stock directo", "—"),
]
tbl_iv = doc.add_table(rows=len(iv_tabla), cols=4)
tbl_iv.style = 'Table Grid'
for i, row_data in enumerate(iv_tabla):
    r = tbl_iv.rows[i].cells
    for j, val in enumerate(row_data):
        r[j].text = val
        run = r[j].paragraphs[0].runs[0] if r[j].paragraphs[0].runs else r[j].paragraphs[0].add_run(val)
        run.font.size = Pt(9)
        if i == 0:
            run.bold = True
            shade_cell(r[j], "1F497D")
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
for row in tbl_iv.rows:
    for j, w in enumerate([4.0, 2.5, 4.5, 5.0]):
        row.cells[j].width = Cm(w)

doc.add_paragraph()

# ─── Glosario rápido de estrategias ──────────────────────────────────────────
add_heading(doc, "Glosario rápido de estrategias mencionadas", level=2)

estrategias_glosario = [
    ("Long Call",
     "Comprar una opción call. Perfil: beneficio ilimitado, pérdida máxima = prima pagada. "
     "Requiere movimiento alcista suficiente antes del vencimiento. Mejor con IV baja."),
    ("Long Put",
     "Comprar una opción put. Mirror bajista del Long Call. Mejor con IV baja."),
    ("Bull Call Spread",
     "Comprar call ATM + vender call OTM (al precio target). Reduce coste y IV-risk. "
     "Beneficio máximo = diferencia de strikes − prima neta. Ideal IV 30–60%."),
    ("Bear Put Spread",
     "Comprar put ATM + vender put OTM bajista (al precio target). Mirror bajista del Bull Call Spread."),
    ("Call Credit Spread (Bear)",
     "Vender call OTM + comprar call más OTM. Estrategia bajista: cobra prima, gana si el precio "
     "no sube. Ideal para parabólicos con IV muy alta (>60%) porque la IV alta hace que la prima "
     "cobrada sea grande y el trade sea favorable incluso con poco movimiento."),
    ("Put Credit Spread (Bull)",
     "Vender put OTM + comprar put más OTM. Alcista: gana si el precio no baja. "
     "Útil cuando la acción tiene alto put skew (miedo de caída exagerado) y se espera rebote."),
    ("Long Call Diagonal",
     "Comprar call de vencimiento lejano (ATM) + vender call de vencimiento cercano (OTM). "
     "Reduce coste de la opción larga recogiendo theta del spread corto. "
     "Ideal para setups con movimiento gradual que necesitan 60–90 días."),
    ("Long Put Diagonal",
     "Mirror bajista del Long Call Diagonal. Sell near-term OTM put + buy far-term ATM put."),
    ("Poor Man's Covered Call (PMCC)",
     "Comprar LEAPS call profundamente ITM (delta ~0.80) + vender calls OTM cercanas mensualmente. "
     "Actúa como poseer el subyacente pero con una fracción del capital. "
     "Ideal para setups Omega LARGO con CrossDW muy fresco que sugiere tendencia de meses."),
    ("Broken Wing Butterfly (BWB)",
     "Butterfly asimétrica: se construye para coste 0 o crédito pequeño. "
     "Ejemplo bajista: comprar 1 put ATM, vender 2 puts OTM-1, comprar 1 put OTM-2 desplazado. "
     "Máximo beneficio en zona media; pérdida controlada. "
     "Ideal para stocks en distribución con target muy claro (Distribución Koncorde)."),
    ("Ratio Spread",
     "Vender más opciones OTM de las que se compran ATM. Genera crédito o bajo coste, "
     "pero expone a pérdida ilimitada en el lado vendido si el movimiento es extremo. "
     "Solo para operadores avanzados con gestión activa."),
]

for (nombre, desc) in estrategias_glosario:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    r = p.add_run(f"{nombre}: ")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.add_run(desc).font.size = Pt(9)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 1. FAMÍLIA KONCORDE
# ════════════════════════════════════════════════════════════════════
add_heading(doc, "Opciones por Screener — Familia Koncorde", level=2)

# ── 1.1 Acumulación ──────────────────────────────────────────────
add_heading(doc, "Koncorde Acumulación v2.1  [LARGO]", level=3)
add_strategy_table(doc, [
    ("IV esperada",          "BAJA (20–40%). La acumulación institucional es silenciosa: sin noticias públicas, sin gap, sin volumen de titular. El mercado no 've' lo que está pasando → IV deprimida → opciones baratas."),
    ("Ventana temporal",     "El setup puede tardar 4–10 semanas en resolver. El azul cruzando cero (Primavera) es el detonante. Comprar tiempo suficiente es crítico."),
    ("Estrategia principal", "LONG CALL ATM · Vencimiento: 60–90 días\nRazón: IV baja hace que la opción sea barata. Se compra antes de que el mercado descubra el movimiento (IV sube al aflorar noticias). La ganancia viene del delta + vega (expansión de IV)."),
    ("Estrategia alternativa","LONG CALL DIAGONAL (60-90d comprado ATM / 30d vendido OTM+10%)\nSi el target es claro y se quiere reducir coste, vender la call cercana cada mes amortiza la posición larga mientras la acumulación sigue en marcha."),
    ("Strike recomendado",   "ATM o máximo 1 strike OTM. El movimiento será grande pero no inmediato."),
    ("Vencimiento óptimo",   "60–90 días desde la entrada. Renovar si se acerca a 21 días sin haber resuelto."),
    ("Gestión",              "Salir si azul pierde la condición azul>verde o si el score cae bajo 50% del umbral de entrada. Tomar 50% de beneficios cuando el precio llega a la nube (Ichimoku)."),
    ("Alerta IV",            "Antes de entrar revisar IV Rank: si >60%, cambiar a Bull Call Spread (comprar ATM, vender OTM al target de la Primavera)."),
])

section_divider(doc)

# ── 1.2 Primavera Naciente ───────────────────────────────────────
add_heading(doc, "Koncorde Primavera Naciente v2.0  [LARGO]", level=3)
add_strategy_table(doc, [
    ("IV esperada",          "BAJA–MODERADA (25–50%). El azul acaba de cruzar y el precio ya está sobre la nube. No hay noticia pública todavía → IV aún manejable. Pero el setup es más maduro que Acumulación."),
    ("Ventana temporal",     "El movimiento empieza en 1–3 semanas desde la señal y puede durar 4–8 semanas hasta target +20%."),
    ("Estrategia principal", "LONG CALL ATM · Vencimiento: 45–60 días\nSetup más inmediato que Acumulación → menos tiempo necesario. bsAzul ≤ 3 indica breakout muy próximo."),
    ("Estrategia alternativa","BULL CALL SPREAD si IV > 40%\nComprar call ATM + vender call al target (+20%). Reduce prima pagada en 40–50%, a cambio de cap de beneficio en el target."),
    ("Strike recomendado",   "ATM (delta ~0.50). Si el movimiento ya empezó (precio sube hoy), entrar 1 strike OTM para pagar menos."),
    ("Vencimiento óptimo",   "45–60 días. Con bsAzul=0 (cruce hoy) → 45 días. Con bsAzul=3 → 60 días."),
    ("Gestión",              "Stop conceptual: si azul vuelve a negativo en < 5 barras, el cruce fue falso. Salir con pérdida controlada. Target: precio +20% o pérdida de kumoTop como soporte."),
    ("Edge adicional",       "La confluencia Ichimoku completa (I1+I2+I3+I4) es rarísima. Cuando aparece con Primavera Naciente, el breakout suele ser potente. Incrementar tamaño."),
])

section_divider(doc)

# ── 1.3 Acumulación+Primavera Combinado ──────────────────────────
add_heading(doc, "Koncorde Acumulación + Primavera Combinado v1.0  [LARGO]", level=3)
add_strategy_table(doc, [
    ("IV esperada",          "Depende de la fase activa: BAJA si acumOK, BAJA-MODERADA si primOK."),
    ("Estrategia según fase","— Si close < Nube (acumOK): Long Call Diagonal 90d/45d. Necesita más tiempo.\n— Si close > Nube (primOK): Long Call ATM 45–60d o Bull Call Spread si IV>40%."),
    ("Vencimiento óptimo",   "acumOK: 75–90 días · primOK: 45–60 días"),
    ("Nota operativa",       "Este screener es útil para monitorear; la estrategia de opciones debe adaptarse al estado concreto de cada acción en el momento de entrada. Siempre verificar la fase antes de elegir la estructura."),
])

section_divider(doc)

# ── 1.4 Distribución ─────────────────────────────────────────────
add_heading(doc, "Koncorde Distribución v1.0  [CORTO / SALIDA]", level=3)
add_strategy_table(doc, [
    ("IV esperada",          "MODERADA–ALTA (40–70%). El verde alto (retail muy comprado) suele coincidir con precio cerca de máximos recientes, lo que eleva la IV. Además el azul cruzando bajista puede coincidir con primeras noticias negativas."),
    ("Ventana temporal",     "La caída desde una distribución con trampa=verde-azul>20 puede ser brusca (2–4 semanas) o gradual (6–10 semanas). Depende de cuándo capitule el retail."),
    ("Estrategia principal", "BEAR PUT SPREAD · Vencimiento: 45–75 días\nComprar put ATM + vender put OTM bajista al target (−15% a −20% del precio actual). La venta del put OTM recorta el coste en un 35–50% y el target queda bien capturado."),
    ("Estrategia alternativa","CALL CREDIT SPREAD (bajista) si IV > 60%\nVender call OTM + comprar call más OTM. Se cobra prima desde el inicio. Gana si el precio se queda plano o cae. No captura toda la caída pero tiene probabilidad alta cuando el stock ya está en distribución y el verde empieza a ceder."),
    ("Strike recomendado",   "Put comprada: ATM (precio actual). Put vendida: −15% a −20% del precio actual (el target de la caída). Diferencia de strikes ≥ $3–5."),
    ("Vencimiento óptimo",   "45–75 días. Si trampa > 30 (situación muy extrema), el pánico puede ser rápido → 30–45 días."),
    ("Gestión",              "Stop: si azul recupera positivo o verde empieza a caer (retail capitulando antes de lo esperado, toma de ganancias). Target: precio −15% o pérdida del soporte de Kijun semanal."),
    ("Caso especial: BWB",   "Broken Wing Butterfly bajista si el target es muy claro (Kijun semanal como soporte bajo): comprar put ATM, vender 2 puts OTM-1 (target), comprar 1 put OTM-2. Puede construirse a coste 0 o incluso crédito pequeño."),
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 2. FAMILIA VIGÍA10
# ════════════════════════════════════════════════════════════════════
add_heading(doc, "Opciones por Screener — Familia Vigía10 / Pre-Breakout", level=2)

add_para(doc,
    "ADVERTENCIA CRÍTICA para esta familia: las acciones $1–$5 (micro-caps y pennies) "
    "tienen IV habitualmente >80–150%. Comprar opciones a esos niveles de IV implica que "
    "incluso si la acción sube 30%, el IV crush post-spike puede eliminar la mayor parte "
    "del beneficio. La regla en esta familia es: siempre usar spreads o considerar "
    "comprar el subyacente directamente.", bold=True, color=(0xC0, 0x00, 0x00))

section_divider(doc)

# ── 2.1 Pre-Breakout v2.0 ────────────────────────────────────────
add_heading(doc, "Vigía10 Pre-Breakout v2.0  [LARGO]", level=3)
add_strategy_table(doc, [
    ("IV esperada",          "ALTA (60–120%). Micro-caps $1–$10: IV elevada por naturaleza."),
    ("Estrategia por tipo de señal", ""),
])

doc.paragraphs[-1].runs[0].font.size = Pt(10)  # ajuste

add_para(doc, "La estrategia varía según el tipo detectado:", bold=True, indent=True)

tipos_opt = [
    ("TipoB (señal más fuerte — SV cruza 0)",
     "Movimiento en 3–10 días. IV alta pero el movimiento es inminente y explosivo.\n"
     "→ Bull Call Spread MUY AJUSTADO: comprar call ATM + vender call OTM+15%. "
     "Vencimiento: 21–30 días. La rapidez del movimiento compensa el IV elevado.\n"
     "Alternativa: Long Call ATM 21 días si IV rank < 50% (poco frecuente en pennies)."),
    ("TipoA (SV negativo pero girando hacia MR)",
     "Señal temprana: necesita 2–4 semanas para que SV cruce cero.\n"
     "→ Bull Call Spread 30–45 días. Más tiempo para que el movimiento se desarrolle."),
    ("TipoC (todo positivo — confirmación tardía)",
     "Parte del movimiento ya ocurrió. IV puede estar ya inflada.\n"
     "→ Bull Call Spread ajustado 21–30 días o PUT CREDIT SPREAD si el precio "
     "ya subió 10%+ (vender el riesgo de retroceso, no perseguir el movimiento)."),
    ("TipoD (divergencia filtro/SV — muy temprano)",
     "El setup más anticipado: necesita paciencia.\n"
     "→ Long Call Diagonal: comprar call 60d ATM + vender call 30d OTM+10%. "
     "Renovar la venta mensualmente hasta que SV cruce cero (TipoB)."),
]

for (tipo, desc) in tipos_opt:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    r = p.add_run(f"{tipo}:  ")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x00, 0x4E, 0x96)
    p.add_run(desc).font.size = Pt(9)

add_strategy_table(doc, [
    ("Vencimiento óptimo",   "TipoB: 21–30 días · TipoA: 30–45 días · TipoC: 21–30 días · TipoD: Diagonal 60/30 días"),
    ("Strike recomendado",   "Siempre ATM en la comprada. OTM vendida en el target (+15% para TipoB, +20% para TipoA/D)."),
    ("Alternativa si no hay opciones líquidas", "Comprar el subyacente directamente con stop por debajo del soporte (kumoBot o mínimo de 5 días)."),
    ("Gestión",              "Si SV vuelve a negativo antes de cruzar 0: stop. Si cruza 0 (TipoB), escalar posición."),
])

section_divider(doc)

# ── 2.2 Absorción Pre-Squeeze v1.5 ──────────────────────────────
add_heading(doc, "Absorción Pre-Squeeze v1.5  [LARGO]", level=3)
add_strategy_table(doc, [
    ("IV esperada",          "MUY ALTA (100–200%). Pennies $0.50–$3 con spikes de volumen."),
    ("Disponibilidad de opciones", "La mayoría de acciones en este rango NO tienen opciones líquidas. Verificar siempre el open interest antes de entrar con opciones."),
    ("Si hay opciones (stock $2–$3)",
     "BULL CALL SPREAD muy ajustado: comprar call ATM + vender call OTM+20%. "
     "Vencimiento: 15–21 días (el squeeze es rápido o no ocurre). "
     "El spread elimina el riesgo de IV crush: si el squeeze ocurre en 3–5 días, "
     "el delta de la call comprada gana más que lo que pierde la call vendida."),
    ("Si hay opciones (stock $1–$2)",
     "Posición muy pequeña en Long Call ATM 15 días. Tamaño: no más del 0.5% del portfolio. "
     "IV crush puede ser devastador si el squeeze no ocurre en la primera semana."),
    ("Estrategia preferida (sin opciones)",
     "Comprar el subyacente directamente: stop < rango del spike de volumen (el soporte de absorción). "
     "Target: EMA21 (+15–25%). Tamaño small dado el riesgo de la clase de activo."),
    ("Vencimiento óptimo",   "15–21 días máximo si se usan opciones."),
    ("Gestión",              "Este es un trade de evento. Si el volumen se seca sin movimiento de precio: salir. Máximo 2 semanas de espera."),
])

section_divider(doc)

# ── 2.3 Sniper-Squeeze v5.3 ─────────────────────────────────────
add_heading(doc, "Vigía10 Sniper-Squeeze v5.3  [LARGO]", level=3)
add_strategy_table(doc, [
    ("IV esperada",          "MUY ALTA (80–200%). Penny stocks $0.10–$5."),
    ("Opciones disponibles", "Raramente hay opciones líquidas en este universo."),
    ("Cuando hay opciones ($3–$5, open interest >500)",
     "BULL CALL SPREAD 21–30 días: ATM comprada + OTM+20% vendida. "
     "Enfoque en señales 'dorada' (SV cruzó MR fresco + filtro > 0) que son las de mayor fiabilidad."),
    ("Estrategia preferida", "Comprar subyacente directamente cuando se disparan las señales 'dorada' + 'sqVigia' simultáneamente. Son los setups con mayor score y mayor confirmación."),
    ("Vencimiento opciones", "21–30 días cuando disponibles."),
    ("Nota",                 "La señal 'Takana' (fusión TK en Ichimoku) dentro de este screener eleva la fiabilidad cuando coincide con 'dorada'. Si ambas coinciden, es el setup prioritario de la familia Vigía10."),
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 3. TAKANA EN KUMO
# ════════════════════════════════════════════════════════════════════
add_heading(doc, "Opciones por Screener — Takana en Kumo v6.0  [LARGO]", level=2)

add_para(doc,
    "Este es el screener donde las opciones tienen el MEJOR ratio riesgo/beneficio de todo el sistema. "
    "La razón: la nube Ichimoku comprime el precio (baja IV real) → la IV implícita también cae → "
    "las opciones están baratas justo antes de un breakout. El edge es estructural.",
    bold=False, italic=True, color=(0x00, 0x60, 0x00))

section_divider(doc)
add_strategy_table(doc, [
    ("IV esperada",          "BAJA (20–45%). La nube comprime el precio y el rango diario. IV baja = opciones baratas = momento ideal para comprar."),
    ("Universo",             "$1–$500 con dolarVol ≥ $500k. Para opciones prácticas: focus en stocks $5+."),
    ("Edge de la estrategia","La Takana indica que el precio está a punto de RESOLVER la compresión de la nube. Comprar opciones antes del breakout es el equivalente a comprar volatilidad barata antes de una explosión de IV. Doble beneficio: delta (movimiento de precio) + vega (expansión de IV)."),
    ("Estrategia principal", "LONG CALL ATM · Vencimiento: 30–45 días\nEl breakout desde la nube suele ser rápido (1–2 semanas) una vez que SV cruza sobre 0. La nube fina (<8%) implica poca resistencia → el movimiento es limpio."),
    ("Strike recomendado",   "ATM (precio dentro de la nube ≈ precio de la call ATM será muy cercano al precio actual). Si kumoTop es visible como resistencia, comprar call con strike = kumoTop."),
    ("Vencimiento óptimo",   "30–45 días. Con postSqueeze activo (compresión de SV previa): 30 días es suficiente. Sin compresión previa: 45 días por seguridad."),
    ("Estrategia alternativa","BULL CALL SPREAD si IV > 40%: comprar call en kumoTop + vender call ATH reciente o +15%. Esto captura el breakout de la nube hacia zona libre."),
    ("Estrategia avanzada",  "LONG STRADDLE si la dirección no está clara (aunque el screener filtra alcistas con svSube): comprar call + put ATM 30d. Beneficio si el breakout de la nube es >8% en cualquier dirección. El coste es bajo porque la IV está comprimida."),
    ("Gestión",              "Si el precio regresa bajo kumoBot (sale de la nube por abajo): stop. El setup fracasó. Si rompe kumoTop con volumen (s13 = ruptura 10 días): escalar. Target: +15–25% desde kumoTop."),
    ("Señal de mejor calidad","Score ≥ 100 + s8 (kumoAlcista) + s13 (ruptura 10 días) + s14 (postSqueeze) = setup de máxima calidad. En este caso preferir Long Call sobre Spread para capturar el movimiento completo."),
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 4. FAMILIA IMÁN KIJUN
# ════════════════════════════════════════════════════════════════════
add_heading(doc, "Opciones por Screener — Familia Imán Kijun", level=2)

# ── 4.1 Imán Kijun v2.0 Largo ────────────────────────────────────
add_heading(doc, "Imán Kijun v2.0  [LARGO]", level=3)
add_strategy_table(doc, [
    ("IV esperada",          "MODERADA (30–60%). Stocks $0.50–$25 en pullback. IV elevada por la caída reciente pero sin pánico extremo."),
    ("Filosofía de la estrategia", "El kijun plano es el target con alta probabilidad estadística. Esto permite construir spreads con las piernas en los niveles exactos: comprar donde está el precio, vender en el kijun."),
    ("Estrategia principal", "BULL CALL SPREAD · Vencimiento: 21–45 días\nComprar call ATM (precio actual) + vender call en el kijun (target).\nEjemplo: precio $12, kijun $14.80 (23% arriba) → comprar call $12 + vender call $15. Prima neta ≈ 50–60% del spread. Beneficio máximo si llega al kijun."),
    ("Estrategia alternativa","PUT CREDIT SPREAD (alcista) si el put skew está muy elevado:\nVender put OTM (−10% del precio actual) + comprar put más OTM (−15%). Recibe crédito. Gana si el precio se mantiene sobre el strike vendido. Ideal cuando se ve que el mercado 'teme' más caída de la que realmente hay."),
    ("Strike recomendado",   "Call comprada: ATM (precio actual). Call vendida: kijun o primer nivel de resistencia visible (tenkan / kumoBot si está sobre el precio)."),
    ("Vencimiento óptimo",   "distKijun 15–20%: 21–30 días. distKijun 20–30%: 30–45 días. distKijun 30–40%: 45 días."),
    ("Gestión",              "Stop: si distKijun aumenta en vez de reducirse 3 días seguidos (reversión activa falla). Profit parcial al 50% del movimiento al kijun. Profit completo en kijun."),
    ("Nota sobre precio bajo $5", "Para stocks $0.50–$5 con opciones disponibles: usar Long Call 21d ATM si IV < 80%, o directamente comprar el subyacente. Los spreads en precios muy bajos tienen spreads bid/ask muy amplios que destruyen la eficiencia."),
])

section_divider(doc)

# ── 4.2 Imán Kijun SHORT v2.0 ────────────────────────────────────
add_heading(doc, "Imán Kijun SHORT v2.0  [CORTO — Bajo la Nube]", level=3)
add_strategy_table(doc, [
    ("IV esperada",          "MODERADA-ALTA (40–75%). Stocks en downtrend bajo la nube: el mercado está nervioso, la IV del lado put tiene skew elevado. Esto paradójicamente ENCARECE los puts que queremos comprar."),
    ("Filosofía de la estrategia", "El rebote temporal sobre el kijun es el momento óptimo para entrar: la acción subió (SV alto = IV baja momentáneamente por el rebote positivo) pero sigue bajo la nube. La IV baja relativa en el rebote es la ventana para comprar puts más baratos."),
    ("Estrategia principal", "BEAR PUT SPREAD · Vencimiento: 30–60 días\nComprar put ATM (precio actual sobre kijun) + vender put en el kijun (target bajista).\nEjemplo: precio $18 (28% sobre kijun $14) → comprar put $18 + vender put $14. Prima neta reducida por la venta."),
    ("Estrategia con rechazo de nube (s1–s4 activos)", "El Rechazo de Nube es la señal más fuerte del screener (hasta 30 pts). Cuando se activa, la caída puede ser rápida y pronunciada. En este caso: LONG PUT ATM 30 días si IV < 50%. La velocidad del rechazo compensa el coste más alto."),
    ("Estrategia alternativa","CALL CREDIT SPREAD bajista si IV > 60%:\nVender call OTM (justo bajo nubeBot) + comprar call más OTM (nubeTop). La nube actúa de techo perfecto para el strike vendido. Recoge prima aprovechando que la nube como resistencia es robusta."),
    ("Strike recomendado",   "Put comprada: ATM (precio actual). Put vendida: kijun o Kijun −5% (la recuperación natural). Spreads de $3–$8 son los más eficientes."),
    ("Vencimiento óptimo",   "distAbove 20–30%: 30–45 días. distAbove 30–60%: 45–60 días. Rechazo de nube activo: 21–30 días (caída rápida esperada)."),
    ("Orden de prioridad",   "1. Rechazo de nube (s1–s4) activo → Long Put 30d. 2. Rechazo sin s1–s4 → Bear Put Spread 45d. 3. IV > 60% → Call Credit Spread bajo nubeBot."),
    ("AVISO: Short directo", "Solo en stocks con dolarVol >$5M. Para el resto: opciones exclusivamente (put o spread). Evitar short directo en stocks $2–$10 bajo la nube = riesgo de short squeeze."),
    ("Gestión",              "Stop: si precio supera nubeBot (entra en la nube). El setup de 'abajo la nube + rebote' falla si la nube es penetrada. Target: kijun. Profit parcial al 50% del recorrido al kijun."),
])

section_divider(doc)

# ── 4.3 Imán Kijun SHORT Sobre Nube v1.0 ─────────────────────────
add_heading(doc, "Imán Kijun SHORT Sobre Nube v1.0  [CORTO — Sobre la Nube]", level=3)
add_strategy_table(doc, [
    ("IV esperada",          "ALTA (50–120%). Stock parabólico: la subida explosiva genera IV muy elevada. Este es el escenario donde la IV cara AYUDA al corto porque podemos VENDER esa IV cara."),
    ("Filosofía de la estrategia", "La acción subió parabolicamente (80–200% en poco tiempo). La IV implícita está en máximos. El SV está girando a la baja. El kijun es el target obvio. Aquí no se compra prima cara: se vende la IV elevada a favor."),
    ("Estrategia principal (PREMIUM SELLING)", "CALL CREDIT SPREAD (bajista) · Vencimiento: 30–45 días\nVender call OTM (precio actual ×1.05–1.10) + comprar call más OTM (precio actual ×1.20).\nSe cobra prima aprovechando IV alta. Gana si el stock no sigue subiendo. Riesgo máximo = diferencia de strikes − crédito cobrado.\nEjemplo: acción en $50 (parabolica desde $20). IV=90%. Vender call $55 + comprar call $65 = cobrar $2.50. Gana si precio queda bajo $55 a vencimiento."),
    ("Estrategia alternativa principal", "BEAR PUT SPREAD · Vencimiento: 30–45 días\nComprar put ATM + vender put OTM al kijun. Requiere desembolso, pero captura la caída directamente.\nMejor cuando la corrección ya ha comenzado y la IV empieza a bajar (pero sigue alta)."),
    ("Estrategia avanzada", "LONG PUT DIAGONAL · Vencimiento: comprar put 60d ATM + vender put 30d OTM (-15%).\nReducción de coste: la put vendida recauda theta durante el primer mes. Si en 30 días la corrección no alcanzó el target, renovar la put vendida."),
    ("Estrategia máxima precisión", "BROKEN WING BUTTERFLY bajista · Coste 0 o crédito pequeño.\nComprar 1 put ATM + vender 2 puts OTM-1 (kijun) + comprar 1 put OTM-2 (kijun −10%).\nBeneficio máximo exactamente en el kijun. Pérdida controlada en ambos lados."),
    ("Strike recomendado",   "Kijun como pivot central de toda la estructura. Es el target estadístico donde el precio converge."),
    ("Vencimiento óptimo",   "30–45 días (screener recomienda 4–8 semanas). Con s14 (volumen decreciente en cima) activo: 30 días. Con s33 (historial sobre nube 20 días): 45 días, la corrección puede ser lenta."),
    ("Orden de estrategias por probabilidad de éxito",
     "1. Call Credit Spread (IV alta → cobrar prima) — mayor prob.\n"
     "2. Bear Put Spread (captura directa de la caída)\n"
     "3. Broken Wing Butterfly (si el kijun es target muy preciso)\n"
     "4. Long Put Diagonal (si se quiere reducir coste con estructura temporal)"),
    ("Gestión",              "Stop: si precio hace nuevo máximo absoluto (el parabólico continúa, la corrección fracasa). Target: kijun (tomar 70% de beneficio). Si el kijun se mantiene plano y el precio sigue cayendo: dejar correr hasta la nube como segundo target."),
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 5. FÓRMULA OMEGA
# ════════════════════════════════════════════════════════════════════
add_heading(doc, "Opciones por Screener — Fórmula Omega", level=2)

add_para(doc,
    "La Fórmula Omega es el screener con mayor alineación multi-sistema. "
    "Cuando score ≥ 200 + CrossDW fresco, el movimiento esperado (≥20%) tiene alta probabilidad "
    "de ocurrir en un horizonte de 4–10 semanas. Esto permite usar estrategias con mayor "
    "duración y mayor convicción en el dimensionamiento.", italic=True)

section_divider(doc)

# ── 5.1 Omega LARGO ──────────────────────────────────────────────
add_heading(doc, "Fórmula Omega LARGO v1.0  [LARGO]", level=3)
add_strategy_table(doc, [
    ("IV esperada",          "VARIABLE (25–65%). Depende del precio y sector. Stocks $5–$15: IV alta (50–80%). Stocks $30–$200: IV moderada (25–45%)."),
    ("Ventana temporal",     "CrossDW fresco (bsCrossDW ≤ 5) = tendencia semanal recién activada. El movimiento del 20% típicamente ocurre en 4–10 semanas."),
    ("Estrategia principal (stocks $30+)", "LONG CALL ATM · Vencimiento: 60–90 días\nIV moderada + movimiento grande esperado + horizonte amplio = el Long Call es la estructura más eficiente. Captura delta + vega (la IV sube cuando el breakout se confirma públicamente)."),
    ("Estrategia principal (stocks $5–$30)", "BULL CALL SPREAD · Vencimiento: 60–75 días\nIV más alta en este rango. Comprar call ATM + vender call en el target (+20%). Reduce coste en 35–50%. Beneficio máximo = +20% desde entrada."),
    ("Estrategia avanzada — PMCC",
     "POOR MAN'S COVERED CALL · Para setups con bsCrossDW ≤ 5 + score ≥ 230 (máxima convicción):\n"
     "Comprar LEAPS call 6 meses ITM (delta ≥ 0.75) + vender call OTM 30d (OTM +10%).\n"
     "Actúa como el subyacente con 25–30% del capital. Renovar la venta mensualmente.\n"
     "Coste del LEAPS se amortiza con las ventas mensuales. El trade puede durar 3–6 meses "
     "si el CrossDW semanal sigue activo."),
    ("Estrategia avanzada — Diagonal",
     "LONG CALL DIAGONAL · 90d/30d:\n"
     "Comprar call 90d ATM + vender call 30d OTM+10%.\n"
     "Reduce coste inicial en 30–40%. Renovar la venta a los 30d. "
     "Si a los 30d el precio subió +10%, cerrar la vendida y comprar el subyacente o escalar."),
    ("Strike recomendado",   "Call comprada: ATM. Call vendida (si spread): precio actual +20% (el target de la Omega). "
                             "Si score ≥ 240: considerar solo la comprada (Long Call puro) para capturar todo el movimiento."),
    ("Vencimiento óptimo",   "bsCrossDW ≤ 5 (muy fresco): 75–90 días. bsCrossDW 6–15 (reciente): 60–75 días. bsCrossDW 16–25 (válido pero menos fresco): 45–60 días."),
    ("Señal de mayor calidad para opciones",
     "K1 + V1 + DW1 simultáneos (azul cruzó hoy, SV cruzó 0 ayer/hoy, DW muy fresco) + score ≥ 230 = "
     "el setup más potente. En este caso: Long Call ATM 75d sin spread. Tamaño más grande."),
    ("Gestión",              "Stop conceptual: si azul pierde positivo O si SV cruza bajo MR antes de que el precio suba 10%. "
                             "Profit parcial (50%) al +10%. Profit total al +20% o al perder el kijun como soporte."),
    ("Escalada",             "Si a las 3 semanas el precio subió +8%+ y el Omega sigue en condición (score > umbral), "
                             "considerar escalar con un segundo spread o rodar la posición a un strike más alto."),
])

section_divider(doc)

# ── 5.2 Omega CORTO ─────────────────────────────────────────────
add_heading(doc, "Fórmula Omega CORTO v1.0  [CORTO]", level=3)
add_strategy_table(doc, [
    ("IV esperada",          "MODERADA-ALTA (35–70%). Las acciones en distribución suelen tener IV moderada inicialmente, que sube cuando el mercado 'huele' el problema."),
    ("La variable 'Trampa'", "La columna Trampa (verde−azul) es el predictor de intensidad de la caída. Trampa > 30 = distribución extrema = los puts del mercado ya empiezan a estar caros. Trampa < 20 = distribución inicial = IV todavía baja = momento para comprar puts baratos."),
    ("Estrategia según Trampa", ""),
])

add_para(doc, "Estrategia según nivel de Trampa:", bold=True, indent=True)
trampa_cases = [
    ("Trampa 10–20 (distribución incipiente)",
     "IV moderada baja. LONG PUT ATM 60–75 días.\n"
     "La señal es fresca, el mercado no lo ha visto. Puts todavía baratos. "
     "Mayor beneficio posible al capturar la caída completa."),
    ("Trampa 20–30 (distribución activa)",
     "IV moderada. BEAR PUT SPREAD 45–60 días.\n"
     "Comprar put ATM + vender put al target (precio −20%).\n"
     "Equilibrio entre coste y captura del movimiento."),
    ("Trampa > 30 (trampa máxima)",
     "IV alta. La caída puede ser brusca y profunda.\n"
     "Dos opciones:\n"
     "  1. BEAR PUT SPREAD ajustado 30–45 días: mover el spread más OTM para reducir coste.\n"
     "  2. PUT DIAGONAL 60/30 días: comprar put ATM 60d + vender put OTM-15% 30d.\n"
     "     La venta del put cercano recauda mientras se espera la capitulación del retail."),
]
for (caso, desc) in trampa_cases:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    r = p.add_run(f"{caso}:  ")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    p.add_run(desc).font.size = Pt(9)

add_strategy_table(doc, [
    ("Estrategia alternativa universal", "CALL CREDIT SPREAD bajista · Vencimiento: 30–45 días\n"
     "Cuando IV > 50%: vender call OTM (precio actual +5%) + comprar call OTM+15%.\n"
     "La nube actúa de resistencia perfecta para el strike vendido. "
     "No requiere que el precio caiga dramáticamente: solo que no siga subiendo."),
    ("Strike recomendado",   "Put comprada: ATM. Put vendida: precio −20% (el target Omega). "
                             "Call vendida (credit spread): precio actual +5% o nubeTop si está cerca. "
                             "Call comprada: precio actual +15%."),
    ("Vencimiento óptimo",   "bsCrossDW ≤ 5: 60–75 días. bsCrossDW 6–15: 45–60 días. bsCrossDW 16–25: 30–45 días. "
                             "Cuanto más fresco el CrossDW, más tiempo necesita para desarrollarse plenamente."),
    ("Gestión",              "Stop: si azul recupera positivo (institucionales vuelven = distribución abortada).\n"
                             "Profit parcial (50%) al −10%. Profit total al −20% (target Omega) o al perder el kijun como soporte."),
    ("Nota sobre borrowing", "Para short directo: verificar borrowable shares + costo de borrow antes de entrar. "
                             "Si borrow rate > 20% anual: usar puts o spreads exclusivamente. "
                             "El coste de borrow puede comerse el beneficio esperado."),
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 6. PENNY SPRING EXTREME
# ════════════════════════════════════════════════════════════════════
add_heading(doc, "Opciones por Screener — Penny Spring Extreme v1.1  [LARGO]", level=2)

add_strategy_table(doc, [
    ("IV esperada",          "EXTREMA (100–250%). Pennies $0.10–$5 bajo la nube: IV siempre muy alta."),
    ("Disponibilidad de opciones", "Stocks $0.10–$1.00: SIN opciones prácticamente. Stocks $1–$3: raramente opciones líquidas. Stocks $3–$5: posiblemente opciones disponibles con bid/ask amplio."),
    ("Estrategia para stocks $3–$5 con opciones",
     "BULL CALL SPREAD muy ajustado · Vencimiento: 15–21 días\n"
     "La señal es un evento de absorción + salto brusco del SV. El move es rápido (1–2 semanas) o no ocurre.\n"
     "Comprar call ATM + vender call OTM+20% (kumoBot como target inicial).\n"
     "Si spikeX ≥ 5 + svJump ≥ 40 simultáneos: máxima convicción → Long Call 15d ATM puro."),
    ("Estrategia preferida universal",
     "COMPRAR EL SUBYACENTE DIRECTAMENTE.\n"
     "Razón: opciones con IV > 100% tienen un IV crush devastador si el squeeze no ocurre en 3–5 días. "
     "El subyacente captura el squeeze sin riesgo de IV crush. Stop: por debajo del rango de la barra de absorción (el soporte es claro)."),
    ("Edge del svJump ≥ 40",
     "Cuando svJump ≥ 40 (s20 del screener), el Vigía10 detecta una compra institucional súbita masiva. "
     "Si además absorcionQuieta está activo, el precio está absorbido: toda la oferta fue retirada. "
     "En este setup específico: Long Call 15d ATM (si disponible) tiene sentido porque el IV crush "
     "no importa si el squeeze ocurre en < 1 semana con subida del 30%+."),
    ("Targets naturales",    "1er target: EMA21 (+10–20%). 2do target: Kijun (+20–35%). 3er target: kumoBot (+35–60%)."),
    ("Vencimiento opciones", "15–21 días MÁXIMO. Este es un trade de evento, no de tendencia."),
    ("Gestión",              "Si en 5 días hábiles desde la señal el precio no se ha movido +5%: salir. El squeeze no ha ocurrido o hay vendedor institucional más grande que el comprador. Stop técnico: bajo el mínimo de la barra de absorción climática (la barra con 3×+ volumen)."),
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# TABLA RESUMEN OPCIONES
# ════════════════════════════════════════════════════════════════════
add_heading(doc, "Tabla Resumen — Opciones por Screener", level=1)

add_para(doc, "Guía rápida de referencia operativa:", bold=True)

headers_opt = ["Screener", "Dir.", "IV esperada", "Estrategia Principal", "Alternativa (IV alta)", "Vencimiento"]
rows_opt = [
    ("Koncorde Acumulación v2.1",          "L",  "Baja 20–40%",      "Long Call ATM",              "Long Call Diagonal 90/45d",       "60–90 días"),
    ("Koncorde Primavera Naciente v2.0",   "L",  "Baja-Mod 25–50%",  "Long Call ATM",              "Bull Call Spread ATM/+20%",        "45–60 días"),
    ("Koncorde Acum+Prim Combinado v1.0",  "L",  "Baja-Mod",         "Según fase activa",          "—",                               "45–90 días"),
    ("Koncorde Distribución v1.0",         "C",  "Mod-Alta 40–70%",  "Bear Put Spread",            "Call Credit Spread bajista",       "45–75 días"),
    ("Vigía10 Pre-Breakout v2.0",          "L",  "Alta 60–120%",     "Bull Call Spread (por tipo)",  "Long Call Diagonal (TipoD)",    "21–45 días"),
    ("Absorción Pre-Squeeze v1.5",         "L",  "Muy alta >100%",   "Subyacente directo",         "Bull Call Spread 15–21d",          "15–21 días"),
    ("Sniper-Squeeze v5.3",                "L",  "Muy alta >80%",    "Subyacente directo",         "Bull Call Spread 21–30d",          "21–30 días"),
    ("Takana en Kumo v6.0",                "L",  "BAJA 20–45%",      "Long Call ATM (MEJOR RATIO)","Bull Call Spread si IV>40%",       "30–45 días"),
    ("Imán Kijun v2.0",                    "L",  "Mod 30–60%",       "Bull Call Spread (ATM/kijun)","Put Credit Spread (alcista)",     "21–45 días"),
    ("Imán Kijun SHORT v2.0",              "C",  "Mod-Alta 40–75%",  "Bear Put Spread",            "Call Credit Spread bajo nubeBot",  "30–60 días"),
    ("Imán Kijun SHORT Sobre Nube v1.0",   "C",  "ALTA 50–120%",     "Call Credit Spread (vender IV alta)","Bear Put Spread / BWB",   "30–45 días"),
    ("Fórmula Omega LARGO v1.0",           "L",  "Variable 25–65%",  "Long Call ATM ($30+) / Bull CS ($5–30)","PMCC o Diagonal 90/30d","60–90 días"),
    ("Fórmula Omega CORTO v1.0",           "C",  "Mod-Alta 35–70%",  "Bear Put Spread / Long Put según trampa","Call Credit Spread bajista","45–75 días"),
    ("Penny Spring Extreme v1.1",          "L",  "Extrema >100%",    "Subyacente directo",         "Bull Call Spread 15–21d",          "15–21 días"),
]

tbl_res = doc.add_table(rows=1, cols=6)
tbl_res.style = 'Table Grid'
hdr = tbl_res.rows[0].cells
for i, h in enumerate(headers_opt):
    hdr[i].text = h
    hdr[i].paragraphs[0].runs[0].bold = True
    hdr[i].paragraphs[0].runs[0].font.size = Pt(8)
    shade_cell(hdr[i], "1F497D")
    hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for row_data in rows_opt:
    r = tbl_res.add_row().cells
    for i, val in enumerate(row_data):
        r[i].text = val
        for p in r[i].paragraphs:
            for run in p.runs:
                run.font.size = Pt(8)
        if i == 1:  # dirección
            if val == "L":
                shade_cell(r[i], "E2EFDA")
            elif val == "C":
                shade_cell(r[i], "FDDBD7")

for row in tbl_res.rows:
    for i, w in enumerate([5.0, 0.7, 2.8, 4.8, 4.5, 2.2]):
        row.cells[i].width = Cm(w)

doc.add_paragraph()
add_para(doc, "L = Largo  ·  C = Corto  ·  CS = Call Spread  ·  BWB = Broken Wing Butterfly  ·  PMCC = Poor Man's Covered Call",
    italic=True, size=8, color=(0x60, 0x60, 0x60))

doc.add_paragraph()

# ─── Nota final de gestión de riesgo ─────────────────────────────
add_heading(doc, "Principios de gestión de riesgo con opciones", level=2)

principios = [
    ("Tamaño de posición",
     "Máximo 2–3% del portfolio por trade en primas pagadas. Para spreads: el riesgo máximo "
     "(diferencia de strikes − crédito/débito) no debe superar el 3% del portfolio."),
    ("Regla del 50% de beneficio",
     "Tomar 50% de la posición cuando el beneficio alcanza el 50% del máximo teórico del spread "
     "o cuando el precio llega a la mitad del target. Deja correr el resto a vencimiento o target."),
    ("Stop del 50% de pérdida",
     "Si la prima pagada cae al 50% de su valor original: salir. No esperar a vencimiento "
     "esperando un milagro. Los milagros en opciones destruyen el portfolio."),
    ("Ajuste por IV Rank",
     "Antes de cualquier entrada, revisar el IV Rank (IVR) del subyacente:\n"
     "IVR < 25%: Long Call/Put (comprar volatilidad barata)\n"
     "IVR 25–50%: Spreads verticales\n"
     "IVR > 50%: Vender premium (Credit Spreads, PMCC, Diagonals)\n"
     "IVR > 75%: Solo credit spreads o diagonals; nunca long options puras."),
    ("Delta objetivo",
     "Para setups de 20%: delta 0.40–0.55 (ATM). Para setups de 30%+: delta 0.35–0.45 (OTM).\n"
     "Evitar delta < 0.25 salvo para setups con muy alta convicción y IV extremadamente baja."),
    ("Días a vencimiento (DTE)",
     "NUNCA mantener opciones compradas con menos de 21 DTE salvo que el movimiento ya esté "
     "en camino (beneficio > 30%). El theta decay se acelera dramáticamente bajo 21 DTE."),
    ("Correlación con el screener",
     "Si el screener saca la acción del resultado (score cae, condición se rompe) "
     "→ cerrar la posición de opciones aunque no haya llegado al stop de pérdida. "
     "El screener es el sistema de gestión: cuando el setup desaparece, el trade termina."),
]

for (titulo, desc) in principios:
    p = doc.add_paragraph()
    r = p.add_run(f"{titulo}:  ")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.add_run(desc).font.size = Pt(10)
    doc.add_paragraph()

# ─── guardar ────────────────────────────────────────────────────────────────
doc.save(docpath)
print(f"Capítulo de opciones añadido. Guardado en: {docpath}")
